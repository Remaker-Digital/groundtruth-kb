#!/usr/bin/env python3
"""
OrbaTech Technical Evaluation Report Generator
Produces a professional DOCX with embedded charts and diagrams.

© 2026 Remaker Digital, a DBA of VanDusen & Palmeter, LLC. All rights reserved.
"""

import os
import io
import math
from datetime import datetime

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.ticker import FuncFormatter
import numpy as np
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────────
# Per S307: discover repo root from script location (scripts/ at repo root).
from pathlib import Path
BASE = str(Path(__file__).resolve().parent.parent)
LOGO_PATH = os.path.join(BASE, "branding", "logo", "PNG", "NEW-BLOCK-LOGO-HORIZONTAL-LIGHT.png")
OUTPUT_PATH = os.path.join(BASE, "OrbaTech-Technical-Evaluation-Report.docx")
CHART_DIR = os.path.join(BASE, "scripts", "_report_charts")
os.makedirs(CHART_DIR, exist_ok=True)

# ── Color Palette ──────────────────────────────────────────────────────────
REMAKER_RED = "#FF3621"
DARK_GRAY = "#2D2D2D"
MID_GRAY = "#666666"
LIGHT_GRAY = "#F5F5F5"
ACCENT_BLUE = "#2196F3"
ACCENT_GREEN = "#4CAF50"
ACCENT_AMBER = "#FFC107"
ACCENT_RED = "#F44336"
ACCENT_PURPLE = "#9C27B0"

# ── Chart Generation ────────────────────────────────────────────────────────

def save_chart(fig, name, dpi=200):
    path = os.path.join(CHART_DIR, f"{name}.png")
    fig.savefig(path, dpi=dpi, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    return path


def chart_maturity_radar():
    """Radar chart: OrbaTech maturity across 8 dimensions."""
    categories = [
        'Tech Stack', 'Architecture', 'Code Quality',
        'Testing', 'CI/CD', 'Security',
        'Documentation', 'Deployment'
    ]
    scores = [8, 6, 4, 2, 1, 3, 3, 2]  # out of 10

    N = len(categories)
    angles = [n / float(N) * 2 * math.pi for n in range(N)]
    angles += angles[:1]
    scores_plot = scores + scores[:1]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    ax.set_theta_offset(math.pi / 2)
    ax.set_theta_direction(-1)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=11, fontweight='bold', color=DARK_GRAY)

    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(['2', '4', '6', '8', '10'], fontsize=9, color=MID_GRAY)
    ax.yaxis.grid(True, color='#E0E0E0', linewidth=0.5)
    ax.xaxis.grid(True, color='#E0E0E0', linewidth=0.5)

    ax.plot(angles, scores_plot, 'o-', linewidth=2.5, color=REMAKER_RED, markersize=8)
    ax.fill(angles, scores_plot, alpha=0.15, color=REMAKER_RED)

    ax.set_title('OrbaTech Engineering Maturity Assessment', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=30)

    return save_chart(fig, 'maturity_radar')


def chart_competitive_comparison():
    """Horizontal bar chart: competitive feature comparison."""
    features = [
        'Contact Management', 'Pipeline/Deals', 'Email Integration',
        'Reporting/Analytics', 'Mobile App', 'API/Integrations',
        'AI/Automation', 'Multi-language', 'Multi-tenant',
        'Customization', 'Calendar/Tasks', 'Security/Compliance'
    ]
    features.reverse()

    # Scores 0-5: 0=None, 1=Basic, 2=Functional, 3=Good, 4=Strong, 5=Enterprise
    data = {
        'OrbaTech':    [3, 3, 2, 2, 0, 1, 0, 0, 2, 2, 3, 2],
        'Salesforce':  [5, 5, 5, 5, 5, 5, 5, 4, 5, 5, 5, 5],
        'HubSpot':     [5, 5, 5, 4, 4, 5, 4, 4, 4, 4, 4, 4],
        'Zoho CRM':    [4, 4, 4, 4, 4, 4, 3, 4, 4, 4, 4, 3],
        'Pipedrive':   [4, 5, 4, 3, 4, 4, 3, 3, 3, 3, 3, 3],
    }
    for k in data:
        data[k].reverse()

    fig, ax = plt.subplots(figsize=(12, 8))
    y_pos = np.arange(len(features))
    bar_height = 0.15
    colors = [REMAKER_RED, '#1A237E', '#FF6F00', '#2E7D32', '#6A1B9A']

    for i, (name, scores) in enumerate(data.items()):
        offset = (i - 2) * bar_height
        bars = ax.barh(y_pos + offset, scores, bar_height, label=name,
                       color=colors[i], alpha=0.85, edgecolor='white', linewidth=0.5)

    ax.set_xlabel('Capability Level', fontsize=11, color=DARK_GRAY)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(features, fontsize=10, color=DARK_GRAY)
    ax.set_xlim(0, 5.5)
    ax.set_xticks([0, 1, 2, 3, 4, 5])
    ax.set_xticklabels(['None', 'Basic', 'Functional', 'Good', 'Strong', 'Enterprise'],
                       fontsize=9)
    ax.legend(loc='lower right', fontsize=10, framealpha=0.9)
    ax.set_title('Feature Capability Comparison', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='x', alpha=0.3)

    return save_chart(fig, 'competitive_comparison')


def chart_azure_costs():
    """Stacked bar chart: Azure cost estimates at scale tiers."""
    tiers = ['10 Tenants\n(Startup)', '100 Tenants\n(Growth)', '1,000 Tenants\n(Scale)']

    # Monthly costs in USD
    costs = {
        'Compute (App Service)':    [50, 300, 1800],
        'SQL Database':             [30, 250, 1500],
        'Storage & CDN':            [5, 30, 200],
        'Identity (Entra ID B2C)':  [0, 50, 500],
        'Monitoring & Logging':     [10, 50, 200],
        'Networking & DNS':         [10, 30, 100],
        'Email Service':            [0, 20, 100],
        'Backup & DR':              [5, 50, 300],
    }

    fig, ax = plt.subplots(figsize=(10, 7))
    x = np.arange(len(tiers))
    width = 0.5
    bottom = np.zeros(len(tiers))

    colors = ['#1565C0', '#2196F3', '#64B5F6', '#90CAF9',
              '#4CAF50', '#81C784', '#FFC107', '#FF7043']

    for i, (label, vals) in enumerate(costs.items()):
        ax.bar(x, vals, width, label=label, bottom=bottom,
               color=colors[i], edgecolor='white', linewidth=0.5)
        bottom += np.array(vals)

    # Total labels on top
    totals = [sum(v[i] for v in costs.values()) for i in range(3)]
    for i, total in enumerate(totals):
        ax.text(i, total + 30, f'${total:,}/mo', ha='center', va='bottom',
                fontsize=12, fontweight='bold', color=DARK_GRAY)

    ax.set_ylabel('Monthly Cost (USD)', fontsize=11, color=DARK_GRAY)
    ax.set_xticks(x)
    ax.set_xticklabels(tiers, fontsize=11, fontweight='bold', color=DARK_GRAY)
    ax.legend(loc='upper left', fontsize=9, framealpha=0.9)
    ax.set_title('Estimated Azure Monthly Costs by Scale Tier', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.yaxis.set_major_formatter(FuncFormatter(lambda x, _: f'${x:,.0f}'))

    return save_chart(fig, 'azure_costs')


def chart_architecture_current():
    """Create architecture diagram showing current OrbaTech architecture."""
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('OrbaTech Current Architecture', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)

    def draw_box(x, y, w, h, label, color, sublabel=None, fontsize=10):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                                        facecolor=color, edgecolor='#333333',
                                        linewidth=1.5, alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.12 if sublabel else 0), label,
                ha='center', va='center', fontsize=fontsize,
                fontweight='bold', color='white' if color not in ['#FFF9C4', '#E3F2FD', '#F3E5F5', '#E8F5E9'] else DARK_GRAY)
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.2, sublabel,
                    ha='center', va='center', fontsize=8,
                    color='white' if color not in ['#FFF9C4', '#E3F2FD', '#F3E5F5', '#E8F5E9'] else MID_GRAY)

    # Presentation layer
    draw_box(0.5, 6.5, 11, 1.2, 'Presentation Layer', '#1565C0',
             'Blazor Server + WebAssembly (Hybrid) | DevExpress UI Components', fontsize=12)

    # API layer
    draw_box(0.5, 5.0, 5, 1.0, 'REST API Controllers', '#2196F3',
             '7 Controllers | Swagger/OpenAPI')
    draw_box(6, 5.0, 5.5, 1.0, 'ASP.NET Core Identity', '#4CAF50',
             'Cookie Auth | 2FA | 4 Roles')

    # Business logic
    draw_box(0.5, 3.5, 3.5, 1.0, 'Service Layer', '#7B1FA2',
             'Per-entity Services')
    draw_box(4.5, 3.5, 3.5, 1.0, 'Repository Layer', '#7B1FA2',
             'EF Core + Dapper')
    draw_box(8.5, 3.5, 3, 1.0, 'Email Service', '#FF6F00',
             'Windows Service')

    # Data layer
    draw_box(0.5, 2.0, 5, 1.0, 'SQL Server', '#1A237E',
             'EF Core 8.0 | 40+ Stored Procedures | 5 Triggers')
    draw_box(6, 2.0, 5.5, 1.0, 'File Storage', '#455A64',
             'Local Filesystem (C:\\Uploads)')

    # Tenant isolation note
    draw_box(0.5, 0.5, 11, 1.0, 'Multi-Tenant Isolation: Application-Level TenantId Filtering (No DB-Level RLS)', '#FFF9C4',
             fontsize=11)

    # Arrows
    arrow_props = dict(arrowstyle='->', color=MID_GRAY, lw=1.5)
    ax.annotate('', xy=(3, 6.5), xytext=(3, 6.0), arrowprops=arrow_props)
    ax.annotate('', xy=(8.75, 6.5), xytext=(8.75, 6.0), arrowprops=arrow_props)
    ax.annotate('', xy=(2.5, 5.0), xytext=(2.5, 4.5), arrowprops=arrow_props)
    ax.annotate('', xy=(6.25, 5.0), xytext=(6.25, 4.5), arrowprops=arrow_props)
    ax.annotate('', xy=(3, 3.5), xytext=(3, 3.0), arrowprops=arrow_props)

    return save_chart(fig, 'architecture_current')


def chart_azure_recommended():
    """Architecture diagram for recommended Azure deployment."""
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title('Recommended Azure Architecture for OrbaTech CRM', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)

    def draw_box(x, y, w, h, label, color, sublabel=None, fontsize=10):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                                        facecolor=color, edgecolor='#333333',
                                        linewidth=1.5, alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.12 if sublabel else 0), label,
                ha='center', va='center', fontsize=fontsize,
                fontweight='bold', color='white' if color not in ['#FFF9C4', '#E3F2FD', '#F3E5F5', '#E8F5E9', '#FBE9E7'] else DARK_GRAY)
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.2, sublabel,
                    ha='center', va='center', fontsize=8,
                    color='white' if color not in ['#FFF9C4', '#E3F2FD', '#F3E5F5', '#E8F5E9', '#FBE9E7'] else MID_GRAY)

    # CDN / Front Door
    draw_box(0.5, 7.5, 13, 1.0, 'Azure Front Door + CDN', '#0D47A1',
             'SSL Termination | WAF | Geographic Routing (CA/US)', fontsize=12)

    # App tier
    draw_box(0.5, 5.8, 6, 1.2, 'Azure App Service (Production)', '#1565C0',
             'P1v3 | Auto-Scale 2-6 | Deployment Slots')
    draw_box(7, 5.8, 6.5, 1.2, 'Azure App Service (Staging)', '#42A5F5',
             'S1 | Deployment Slot Swap')

    # Services
    draw_box(0.5, 4.2, 3, 1.1, 'Azure SQL Database', '#1A237E',
             'S2 (50 DTU) → Elastic Pool')
    draw_box(4, 4.2, 3, 1.1, 'Azure Blob Storage', '#455A64',
             'File Attachments + Backups')
    draw_box(7.5, 4.2, 3, 1.1, 'Azure Key Vault', '#2E7D32',
             'Secrets | API Keys | Certs')
    draw_box(11, 4.2, 2.5, 1.1, 'Entra ID B2C', '#6A1B9A',
             'SSO | MFA')

    # Supporting services
    draw_box(0.5, 2.7, 3, 1.0, 'Azure Monitor', '#E65100',
             'App Insights | Log Analytics')
    draw_box(4, 2.7, 3, 1.0, 'Azure Comm Services', '#00695C',
             'Email Delivery')
    draw_box(7.5, 2.7, 3, 1.0, 'Azure Redis Cache', '#B71C1C',
             'Session | Tenant Config')
    draw_box(11, 2.7, 2.5, 1.0, 'Azure Backup', '#37474F',
             'SQL + Blob PITR')

    # Security layer
    draw_box(0.5, 1.2, 13, 1.0, 'Security: Row-Level Security (RLS) | Managed Identities | Private Endpoints | RBAC',
             '#E8F5E9', fontsize=11)

    # Compliance note
    draw_box(0.5, 0.2, 13, 0.7,
             'Compliance: PIPEDA (CA) | SOC 2 Type II Ready | Canada Central + Canada East Regions',
             '#FBE9E7', fontsize=10)

    return save_chart(fig, 'azure_recommended')


def chart_priority_roadmap():
    """Gantt-style chart for recommended technical priorities."""
    fig, ax = plt.subplots(figsize=(12, 7))

    tasks = [
        ('Security Remediation', 0, 2, ACCENT_RED, 'P0 — Critical'),
        ('CI/CD Pipeline', 1, 3, '#E65100', 'P0 — Critical'),
        ('Automated Testing', 1, 5, '#E65100', 'P0 — Critical'),
        ('Multi-Tenant RLS', 2, 4, ACCENT_AMBER, 'P1 — High'),
        ('Containerization (Docker)', 3, 5, ACCENT_AMBER, 'P1 — High'),
        ('Azure Deployment (IaC)', 4, 7, ACCENT_BLUE, 'P1 — High'),
        ('API Authentication (OAuth)', 3, 5, ACCENT_AMBER, 'P1 — High'),
        ('Localization (EN/FR)', 5, 7, ACCENT_GREEN, 'P2 — Medium'),
        ('Performance Optimization', 6, 8, ACCENT_GREEN, 'P2 — Medium'),
        ('Mobile Responsive', 7, 10, ACCENT_PURPLE, 'P3 — Future'),
        ('AI/Automation Features', 8, 12, ACCENT_PURPLE, 'P3 — Future'),
        ('Agent Red Integration', 9, 12, ACCENT_PURPLE, 'P3 — Future'),
    ]
    tasks.reverse()

    for i, (name, start, end, color, priority) in enumerate(tasks):
        ax.barh(i, end - start, left=start, height=0.6, color=color,
                alpha=0.8, edgecolor='white', linewidth=0.5)
        ax.text(end + 0.2, i, priority, va='center', fontsize=8, color=MID_GRAY)

    ax.set_yticks(range(len(tasks)))
    ax.set_yticklabels([t[0] for t in tasks], fontsize=10, color=DARK_GRAY)
    ax.set_xlabel('Months', fontsize=11, color=DARK_GRAY)
    ax.set_title('Recommended Technical Priority Roadmap', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_xlim(0, 13)
    ax.grid(axis='x', alpha=0.3)

    # Legend
    legend_elements = [
        mpatches.Patch(facecolor=ACCENT_RED, label='P0 — Critical (Months 1-3)'),
        mpatches.Patch(facecolor=ACCENT_AMBER, label='P1 — High (Months 2-6)'),
        mpatches.Patch(facecolor=ACCENT_GREEN, label='P2 — Medium (Months 5-8)'),
        mpatches.Patch(facecolor=ACCENT_PURPLE, label='P3 — Future (Months 7-12)'),
    ]
    ax.legend(handles=legend_elements, loc='lower right', fontsize=9)

    return save_chart(fig, 'priority_roadmap')


def chart_security_heatmap():
    """Heatmap for security posture."""
    categories = [
        'Authentication', 'Authorization', 'Data Encryption',
        'Secrets Mgmt', 'Input Validation', 'API Security',
        'Tenant Isolation', 'Audit Logging', 'HTTPS/TLS',
        'Dependency Scanning'
    ]
    # 1=Critical Gap, 2=Needs Work, 3=Partial, 4=Adequate, 5=Strong
    scores = [4, 3, 1, 1, 2, 2, 2, 1, 3, 1]

    fig, ax = plt.subplots(figsize=(10, 5))

    colors_map = {1: '#D32F2F', 2: '#F57C00', 3: '#FBC02D', 4: '#66BB6A', 5: '#2E7D32'}
    labels_map = {1: 'Critical Gap', 2: 'Needs Work', 3: 'Partial', 4: 'Adequate', 5: 'Strong'}

    bars = ax.barh(range(len(categories)), scores, color=[colors_map[s] for s in scores],
                   height=0.6, edgecolor='white', linewidth=1)

    for i, (score, cat) in enumerate(zip(scores, categories)):
        ax.text(score + 0.1, i, labels_map[score], va='center', fontsize=9, color=MID_GRAY)

    ax.set_yticks(range(len(categories)))
    ax.set_yticklabels(categories, fontsize=10, color=DARK_GRAY)
    ax.set_xlim(0, 6)
    ax.set_xticks([1, 2, 3, 4, 5])
    ax.set_xticklabels(['Critical\nGap', 'Needs\nWork', 'Partial', 'Adequate', 'Strong'], fontsize=9)
    ax.set_title('Security Posture Assessment', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='x', alpha=0.2)

    return save_chart(fig, 'security_heatmap')


def chart_integration_architecture():
    """Bidirectional integration diagram: OrbaTech <-> Agent Red."""
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_title('Bidirectional Integration: OrbaTech CRM ↔ Agent Red AI', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)

    def draw_box(x, y, w, h, label, color, sublabel=None, fontsize=10):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                                        facecolor=color, edgecolor='#333333',
                                        linewidth=1.5, alpha=0.85)
        ax.add_patch(rect)
        cy = y + h/2 + (0.12 if sublabel else 0)
        ax.text(x + w/2, cy, label, ha='center', va='center', fontsize=fontsize,
                fontweight='bold', color='white')
        if sublabel:
            ax.text(x + w/2, cy - 0.3, sublabel, ha='center', va='center',
                    fontsize=8, color='#DDDDDD')

    # OrbaTech side
    draw_box(0.3, 5.0, 5, 1.3, 'OrbaTech CRM', '#1565C0',
             '.NET 8 | Blazor | SQL Server', fontsize=13)
    draw_box(0.3, 3.2, 2.3, 1.3, 'Contacts &\nAccounts', '#1976D2')
    draw_box(3, 3.2, 2.3, 1.3, 'Opportunities\n& Pipeline', '#1976D2')
    draw_box(0.3, 1.5, 2.3, 1.3, 'Activities &\nEmails', '#1976D2')
    draw_box(3, 1.5, 2.3, 1.3, 'Custom\nFields', '#1976D2')

    # Agent Red side
    draw_box(7.7, 5.0, 5, 1.3, 'Agent Red AI', REMAKER_RED,
             'Python | Azure | Cosmos DB', fontsize=13)
    draw_box(7.7, 3.2, 2.3, 1.3, 'AI Customer\nService', '#D84315')
    draw_box(10.4, 3.2, 2.3, 1.3, 'Knowledge\nBase', '#D84315')
    draw_box(7.7, 1.5, 2.3, 1.3, 'Conversation\nHistory', '#D84315')
    draw_box(10.4, 1.5, 2.3, 1.3, 'Quality\nAnalytics', '#D84315')

    # Integration hub
    draw_box(5.6, 3.5, 1.8, 2.5, 'Integration\nHub', '#2E7D32',
             'REST API\nWebhooks\nOAuth 2.0', fontsize=11)

    # Arrows
    arrow_r = dict(arrowstyle='->', color='#1565C0', lw=2.5)
    arrow_l = dict(arrowstyle='->', color=REMAKER_RED, lw=2.5)

    # OrbaTech → Hub
    ax.annotate('', xy=(5.6, 5.2), xytext=(5.3, 5.2), arrowprops=arrow_r)
    # Hub → Agent Red
    ax.annotate('', xy=(7.7, 5.2), xytext=(7.4, 5.2), arrowprops=arrow_r)
    # Agent Red → Hub (return)
    ax.annotate('', xy=(7.4, 4.6), xytext=(7.7, 4.6), arrowprops=arrow_l)
    # Hub → OrbaTech (return)
    ax.annotate('', xy=(5.3, 4.6), xytext=(5.6, 4.6), arrowprops=arrow_l)

    # Flow labels
    ax.text(6.5, 6.5, 'CRM Data → AI Context', fontsize=10, ha='center',
            color='#1565C0', fontweight='bold')
    ax.text(6.5, 0.7, 'AI Insights → CRM Records', fontsize=10, ha='center',
            color=REMAKER_RED, fontweight='bold')

    return save_chart(fig, 'integration_architecture')


def chart_annex_comparison():
    """Side-by-side comparison table as a visual chart."""
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.axis('off')
    ax.set_title('Annex A: OrbaTech Recommended vs Agent Red Azure Deployment',
                 fontsize=14, fontweight='bold', color=DARK_GRAY, pad=15)

    columns = ['Component', 'OrbaTech\n(Recommended)', 'Agent Red\n(Current)']
    rows = [
        ['Compute', 'App Service P1v3\nAuto-scale 2-6', 'Container Apps\nMin 2 replicas'],
        ['Database', 'Azure SQL\nElastic Pool (DTU)', 'Cosmos DB\nServerless + Provisioned'],
        ['Identity', 'Entra ID B2C\nCookie + OAuth', 'Custom Identity\nWidget Keys + API Keys'],
        ['File Storage', 'Azure Blob\nStandard LRS', 'Azure Blob\nStandard LRS'],
        ['Secrets', 'Azure Key Vault\nManaged Identity', 'Azure Key Vault\nManaged Identity + KEK/DEK'],
        ['CDN/Gateway', 'Azure Front Door', 'Container Apps\nIngress'],
        ['Email', 'Azure Comm Services', 'Azure Comm Services\n+ SMS OTP'],
        ['Monitoring', 'App Insights +\nLog Analytics', 'App Insights +\nCustom Observability'],
        ['CI/CD', 'GitHub Actions\n(To be built)', 'GitHub Actions\n13-phase pipeline'],
        ['Multi-Tenant', 'App-Level +\nSQL RLS', 'App-Level +\nCosmos Partition Keys'],
    ]

    table = ax.table(cellText=rows, colLabels=columns,
                     cellLoc='center', loc='center',
                     colWidths=[0.25, 0.35, 0.35])
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 2.2)

    # Style header
    for j in range(3):
        cell = table[0, j]
        cell.set_facecolor('#1565C0')
        cell.set_text_props(color='white', fontweight='bold', fontsize=10)
        cell.set_edgecolor('white')

    # Style rows
    for i in range(1, len(rows) + 1):
        for j in range(3):
            cell = table[i, j]
            cell.set_facecolor('#F5F5F5' if i % 2 == 0 else 'white')
            cell.set_edgecolor('#E0E0E0')
            if j == 0:
                cell.set_text_props(fontweight='bold')

    return save_chart(fig, 'annex_comparison')


# ── DOCX Generation ────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading_elm)


def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return heading


def add_key_finding_box(doc, title, text, severity='info'):
    """Add a highlighted finding box."""
    colors = {
        'critical': 'FF3621',
        'high': 'F57C00',
        'medium': 'FBC02D',
        'info': '2196F3',
        'positive': '4CAF50',
    }
    color = colors.get(severity, '2196F3')

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F5F5F5')

    p_title = cell.paragraphs[0]
    run = p_title.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)

    p_text = cell.add_paragraph(text)
    p_text.style = doc.styles['Normal']
    for run in p_text.runs:
        run.font.size = Pt(10)

    # Add spacing after
    doc.add_paragraph()


def build_report():
    """Build the full DOCX report."""
    print("Generating charts...")
    chart_paths = {
        'maturity': chart_maturity_radar(),
        'competitive': chart_competitive_comparison(),
        'azure_costs': chart_azure_costs(),
        'architecture': chart_architecture_current(),
        'azure_recommended': chart_azure_recommended(),
        'roadmap': chart_priority_roadmap(),
        'security': chart_security_heatmap(),
        'integration': chart_integration_architecture(),
        'annex_comparison': chart_annex_comparison(),
    }
    print("Charts generated.")

    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('OrbaTech CRM')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Technical Evaluation Report')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xFF, 0x36, 0x21)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Prepared by Remaker Digital')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f'April 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = classification.add_run('CONFIDENTIAL')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xFF, 0x36, 0x21)
    run.bold = True

    doc.add_page_break()

    # ── Footer with logo ──
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add logo to footer
        run = footer_para.add_run()
        # Height of ~2 lines of footer text ≈ 0.35 inches
        run.add_picture(LOGO_PATH, height=Inches(0.35))

    # ── Table of Contents placeholder ──
    add_styled_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '3. Technology Assessment',
        '4. Engineering Maturity',
        '5. Security Assessment',
        '6. Multi-Tenant Architecture',
        '7. Competitive Landscape',
        '8. Reliability & Serviceability',
        '9. Maintainability',
        '10. Performance Considerations',
        '11. Recommended Azure Configuration',
        '12. Azure Cost Projections',
        '13. Recommended Technical Priorities',
        'Annex A: OrbaTech vs Agent Red Azure Deployment',
        'Annex B: OrbaTech + Agent Red Integration',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '1. Executive Summary', level=1)

    doc.add_paragraph(
        'This report provides an independent technical evaluation of the OrbaTech CRM platform, '
        'a Blazor-based customer relationship management system built on Microsoft .NET 8. '
        'The evaluation was conducted by Remaker Digital as part of a partnership assessment.'
    )

    doc.add_paragraph(
        'OrbaTech is an early-stage product (approximately 3.5 months in development) built by a small '
        'team of three contributors. The platform targets small-to-medium businesses in the Canadian and '
        'US markets with core CRM functionality including contact management, deal pipeline tracking, '
        'email integration, and scheduling.'
    )

    add_key_finding_box(doc,
        'Overall Assessment',
        'OrbaTech demonstrates solid technology choices (modern .NET 8 + Blazor) and a functional '
        'core architecture. However, significant gaps exist in testing, CI/CD, security hygiene, '
        'and cloud deployment readiness that must be addressed before commercial production use. '
        'The platform has strong potential but requires structured engineering investment.',
        'info'
    )

    # Maturity radar
    doc.add_picture(chart_paths['maturity'], width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run('Figure 1: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Engineering maturity across eight dimensions (scale of 1-10). '
                     'Tech Stack scores highest; CI/CD and Deployment require immediate attention.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 2. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '2. Project Overview', level=1)

    # Project facts table
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    facts = [
        ('Repository', 'github.com/QGolle/OrbaTech'),
        ('Description', 'Blazor-based CRM for sales team deal management'),
        ('Primary Language', 'C# (.NET 8.0)'),
        ('Project Age', '~3.5 months (since December 2025)'),
        ('Total Commits', '183+'),
        ('Contributors', '3 (1 primary, 2 supporting)'),
        ('License', 'None (all rights reserved)'),
        ('Target Markets', 'Canada and United States (EN primary, FR secondary, ES planned)'),
        ('Website', 'orbatechcrm.com'),
    ]
    for i, (label, value) in enumerate(facts):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        for cell in table.row_cells(i):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if cell == table.cell(i, 0):
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'The project demonstrates an active development cadence with daily commits from the primary '
        'contributor. The codebase comprises approximately 1.5 million bytes of source code across '
        'C#, HTML (Razor), CSS, JavaScript, and T-SQL.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 3. TECHNOLOGY ASSESSMENT
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '3. Technology Assessment', level=1)

    add_key_finding_box(doc,
        'Technology Choices: Strong Foundation',
        'The selection of .NET 8 with Blazor represents a modern, well-supported technology stack '
        'suitable for enterprise CRM applications. Microsoft provides long-term support for .NET 8 '
        'through November 2026, with a clear upgrade path to future LTS releases.',
        'positive'
    )

    # Architecture diagram
    doc.add_picture(chart_paths['architecture'], width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run('Figure 2: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Current OrbaTech application architecture showing the N-tier layout with '
                     'Blazor hybrid rendering, dual ORM strategy, and application-level tenant isolation.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_styled_heading(doc, '3.1 Technology Stack Details', level=2)

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Shading Accent 1'
    headers = ['Layer', 'Technology', 'Assessment']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    stack_rows = [
        ('Frontend', 'Blazor Server + WebAssembly\n(Hybrid Rendering)', 'Modern choice enabling SSR + interactive components. DevExpress UI suite provides enterprise-grade components.'),
        ('Backend', 'ASP.NET Core 8.0', 'Industry-standard, high-performance web framework with excellent tooling and community support.'),
        ('ORM', 'Entity Framework Core 8.0\n+ Dapper 2.1', 'Dual ORM strategy is pragmatic: EF Core for CRUD operations, Dapper for stored procedure calls and performance-critical paths.'),
        ('Database', 'SQL Server', 'Appropriate for structured CRM data. Stored procedures (40+) and triggers (5) indicate investment in data-layer logic.'),
        ('Authentication', 'ASP.NET Core Identity', 'Built-in framework with 2FA support. Adequate for current needs; should migrate to Entra ID B2C for cloud deployment.'),
        ('UI Components', 'DevExpress Blazor 25.2', 'Commercial component suite. Provides grid, dashboard, scheduler, and rich text editor. Reduces development time but adds licensing costs.'),
        ('Email', 'MailKit (Windows Service)', 'Standalone email processing service. Functions as spool watcher. Should be migrated to Azure Communication Services for cloud.'),
    ]
    for i, (layer, tech, assessment) in enumerate(stack_rows):
        table.cell(i+1, 0).text = layer
        table.cell(i+1, 1).text = tech
        table.cell(i+1, 2).text = assessment
        for cell in table.row_cells(i+1):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 4. ENGINEERING MATURITY
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '4. Engineering Maturity', level=1)

    doc.add_paragraph(
        'Engineering maturity was assessed across eight dimensions using a 1-10 scale. '
        'This assessment reflects the current state of the codebase and development practices '
        'as observed in the public repository.'
    )

    # Summary table
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Shading Accent 1'
    table.cell(0, 0).text = 'Dimension'
    table.cell(0, 1).text = 'Score'
    table.cell(0, 2).text = 'Key Observations'
    for c in range(3):
        for run in table.cell(0, c).paragraphs[0].runs:
            run.bold = True

    maturity_rows = [
        ('Tech Stack', '8/10', 'Modern .NET 8 + Blazor. Well-chosen dependencies. Current patch versions.'),
        ('Architecture', '6/10', 'Clean N-tier with repository/service pattern. Good separation of concerns.'),
        ('Code Quality', '4/10', 'No linting tools, no static analysis. Inconsistent nullable annotations. Commented-out code present.'),
        ('Testing', '2/10', 'Test project referenced but not visible in repo. No evidence of regular test execution.'),
        ('CI/CD', '1/10', 'Only a csproj sanitization workflow. No build, test, or deploy automation.'),
        ('Security', '3/10', 'Identity/2FA implemented. However, credentials committed to repo. API keys in plaintext.'),
        ('Documentation', '3/10', 'Developer-oriented README. No architecture docs, no API guide, no deployment instructions.'),
        ('Deployment', '2/10', 'No containerization. No IaC. No cloud deployment configuration.'),
    ]
    for i, (dim, score, obs) in enumerate(maturity_rows):
        table.cell(i+1, 0).text = dim
        table.cell(i+1, 1).text = score
        table.cell(i+1, 2).text = obs
        for cell in table.row_cells(i+1):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    add_styled_heading(doc, '4.1 Coding Standards', level=2)

    doc.add_paragraph(
        'The codebase follows a consistent repository-service pattern with dependency injection, '
        'which demonstrates sound architectural thinking. However, several coding standard gaps were identified:'
    )

    standards = [
        'No .editorconfig file for consistent code formatting across contributors',
        'No Roslyn analyzers or StyleCop for compile-time code quality enforcement',
        'Inconsistent use of C# nullable reference type annotations',
        'Commented-out code in production files (e.g., HTTPS redirection)',
        'Commit messages contain spelling inconsistencies',
        'No code review process visible (no pull request workflow observed)',
    ]
    for item in standards:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 5. SECURITY ASSESSMENT
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '5. Security Assessment', level=1)

    add_key_finding_box(doc,
        'CRITICAL: Credentials Committed to Repository',
        'Connection strings containing server names, a hardcoded API access token, and default '
        'administrator credentials were found committed to the repository. These must be removed '
        'immediately and rotated. All secrets should be managed via Azure Key Vault or .NET User Secrets.',
        'critical'
    )

    doc.add_picture(chart_paths['security'], width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run('Figure 3: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Security posture assessment across ten domains. Red and orange areas require immediate remediation.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_styled_heading(doc, '5.1 Findings Detail', level=2)

    findings = [
        ('CRITICAL', 'Hardcoded connection strings in appsettings.json and appsettings.Production.json expose database server names. An API access token is committed in EmailService/appsettings.json.'),
        ('CRITICAL', 'Default administrator account seeded with the password "Password123!" — while gated to development environment, the seed logic and credentials are visible in the repository.'),
        ('HIGH', 'Database backup files (.bak) committed to the repository may contain production data and should be removed from version control history.'),
        ('HIGH', 'File attachment paths are hardcoded to Windows-specific local directories with no visible path traversal protection.'),
        ('MEDIUM', 'No API rate limiting observed. API key validation exists but keys appear stored in plaintext in the database.'),
        ('MEDIUM', 'HTTPS redirection is commented out, and AllowedHosts is set to wildcard in all environments.'),
    ]

    table = doc.add_table(rows=len(findings)+1, cols=2)
    table.style = 'Light Shading Accent 1'
    table.cell(0, 0).text = 'Severity'
    table.cell(0, 1).text = 'Finding'
    for c in range(2):
        for run in table.cell(0, c).paragraphs[0].runs:
            run.bold = True

    severity_colors = {'CRITICAL': 'FF3621', 'HIGH': 'F57C00', 'MEDIUM': 'FBC02D'}
    for i, (sev, finding) in enumerate(findings):
        table.cell(i+1, 0).text = sev
        table.cell(i+1, 1).text = finding
        # Color the severity cell
        for run in table.cell(i+1, 0).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(severity_colors[sev])
            run.font.size = Pt(9)
        for run in table.cell(i+1, 1).paragraphs[0].runs:
            run.font.size = Pt(9)

    doc.add_paragraph()

    add_styled_heading(doc, '5.2 Multi-Tenant Security Considerations', level=2)
    doc.add_paragraph(
        'In a multi-tenant CRM environment, data isolation between tenants is a fundamental '
        'security requirement. OrbaTech currently relies exclusively on application-level filtering '
        'using a TenantId column on each entity, with no database-level enforcement.'
    )

    doc.add_paragraph(
        'This approach has a known vulnerability: a single missing TenantId filter in any repository '
        'query could expose one tenant\'s data to another. The recommended mitigation is to implement '
        'SQL Server Row-Level Security (RLS) policies as a defense-in-depth measure.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 6. MULTI-TENANT ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '6. Multi-Tenant Architecture', level=1)

    doc.add_paragraph(
        'OrbaTech implements a shared-database, shared-schema multi-tenancy model. '
        'All tenants share a single SQL Server database with tenant isolation enforced through '
        'a TenantId foreign key on every business entity.'
    )

    add_styled_heading(doc, '6.1 Current Implementation', level=2)

    mt_strengths = [
        'TenantId consistently present on all major entities (Account, Contact, Opportunity, etc.)',
        'Tenant claims injected at authentication time via custom ClaimsPrincipalFactory',
        'License tracking (LicenseCount, LicenseUsage) per tenant',
        'Dedicated test data with tenant-identifying names for leak detection',
        'SecurityContext service centralizes tenant resolution',
    ]
    for item in mt_strengths:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, '6.2 Gaps and Recommendations', level=2)

    mt_gaps = [
        ('No EF Core Global Query Filters', 'Add HasQueryFilter on every entity to automatically apply TenantId filtering, preventing accidental data leakage from queries that omit the filter.'),
        ('No Database-Level RLS', 'Implement SQL Server Row-Level Security policies as a defense-in-depth mechanism that prevents cross-tenant access even if application code is bypassed.'),
        ('Integer-Based TenantId', 'Consider migrating to GUID-based tenant identifiers to prevent enumeration attacks and simplify multi-region database merges.'),
        ('No Tenant Provisioning Automation', 'Implement automated tenant onboarding with configuration validation, license enforcement, and audit logging.'),
    ]

    table = doc.add_table(rows=len(mt_gaps)+1, cols=2)
    table.style = 'Light Shading Accent 1'
    table.cell(0, 0).text = 'Gap'
    table.cell(0, 1).text = 'Recommendation'
    for c in range(2):
        for run in table.cell(0, c).paragraphs[0].runs:
            run.bold = True

    for i, (gap, rec) in enumerate(mt_gaps):
        table.cell(i+1, 0).text = gap
        table.cell(i+1, 1).text = rec
        for run in table.cell(i+1, 0).paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        for run in table.cell(i+1, 1).paragraphs[0].runs:
            run.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 7. COMPETITIVE LANDSCAPE
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '7. Competitive Landscape', level=1)

    doc.add_paragraph(
        'OrbaTech enters a mature and competitive CRM market. The following comparison positions '
        'OrbaTech against established players serving the SMB segment in North America.'
    )

    doc.add_picture(chart_paths['competitive'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run('Figure 4: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Feature capability comparison across twelve CRM dimensions. '
                     'OrbaTech shows strength in core CRM functions but has gaps in mobile, API, AI, and localization.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_styled_heading(doc, '7.1 Competitor Profiles', level=2)

    competitors = [
        ('Salesforce', '$25-$300/user/mo', 'Market leader with the most comprehensive feature set. Often perceived as complex and expensive for SMBs. Strong enterprise adoption.'),
        ('HubSpot CRM', 'Free-$1,200/mo', 'Leading SMB CRM with a generous free tier. Strong marketing automation and content management integration. Growing enterprise presence.'),
        ('Zoho CRM', '$14-$52/user/mo', 'Cost-effective alternative with broad feature coverage. Multi-language support (28 languages). Strong in price-sensitive SMB segments.'),
        ('Pipedrive', '$14-$99/user/mo', 'Sales-focused CRM known for intuitive pipeline management. Popular with small sales teams. Limited customization compared to larger platforms.'),
    ]

    table = doc.add_table(rows=len(competitors)+1, cols=3)
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Competitor', 'Pricing', 'Market Position']):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    for i, (name, price, position) in enumerate(competitors):
        table.cell(i+1, 0).text = name
        table.cell(i+1, 1).text = price
        table.cell(i+1, 2).text = position
        for cell in table.row_cells(i+1):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    add_styled_heading(doc, '7.2 OrbaTech Differentiation Opportunities', level=2)

    diff_items = [
        'Canadian data residency (compliance with PIPEDA) — a differentiator for Canadian businesses',
        'Bilingual (EN/FR) native support — required for Quebec market, underserved by most US-based CRMs',
        'Simpler pricing and onboarding for micro-businesses (1-10 users)',
        'Integration with AI customer service (Agent Red partnership potential)',
        'Industry-specific customization for Canadian market verticals',
    ]
    for item in diff_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 8. RELIABILITY & SERVICEABILITY
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '8. Reliability & Serviceability', level=1)

    add_styled_heading(doc, '8.1 Reliability', level=2)

    doc.add_paragraph(
        'Reliability measures the system\'s ability to operate without failure. '
        'The current state presents several reliability concerns:'
    )

    reliability_items = [
        ('No automated testing in CI', 'Without automated test execution, regressions can be introduced with every commit and may not be detected until user-facing failures occur.'),
        ('No health check endpoints', 'No visible health or readiness probe endpoints for load balancer or orchestrator integration.'),
        ('Single-server deployment model', 'No horizontal scaling, load balancing, or failover configuration observed.'),
        ('Windows service dependency', 'The EmailService runs as a Windows service tied to a specific machine, creating a single point of failure.'),
        ('No error tracking', 'No integration with error tracking services (e.g., Application Insights, Sentry) was observed.'),
    ]

    for title, desc in reliability_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    add_styled_heading(doc, '8.2 Serviceability', level=2)

    doc.add_paragraph(
        'Serviceability measures how easily the system can be maintained, updated, and diagnosed in production.'
    )

    service_items = [
        ('Deployment process', 'No documented deployment process. No deployment automation. Updates likely require manual intervention on the server.'),
        ('Logging', 'Standard ASP.NET Core logging configured. No structured logging or centralized log aggregation observed.'),
        ('Diagnostics', 'No diagnostic endpoints, no performance counters, no distributed tracing.'),
        ('Database management', 'EF Core migrations present but no automated migration strategy for zero-downtime deployments.'),
        ('Configuration management', 'Environment-specific appsettings files exist but contain hardcoded values that should be externalized.'),
    ]

    for title, desc in service_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 9. MAINTAINABILITY
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '9. Maintainability', level=1)

    doc.add_paragraph(
        'Maintainability assesses how easily the codebase can be understood, modified, '
        'and extended by current and future developers.'
    )

    add_styled_heading(doc, '9.1 Strengths', level=2)

    strengths = [
        'Consistent repository-service pattern with interface-based dependency injection',
        'Clear project structure with logical separation of concerns (Models, Repositories, Services, Components)',
        'Shared class library for DTOs and interfaces promotes code reuse between server and client projects',
        'Entity Framework Core migrations provide version-controlled schema management',
        'Scoped CSS styles per component prevent style conflicts',
    ]
    for item in strengths:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, '9.2 Concerns', level=2)

    concerns = [
        'Dual ORM strategy (EF Core + Dapper + stored procedures) creates three places where data access logic can live, increasing cognitive load',
        'No automated code quality enforcement — relies entirely on developer discipline',
        'Database backup files and migration archives committed to the repository inflate repository size',
        'No contribution guidelines or coding standards document for onboarding new developers',
        'DevExpress commercial licensing creates vendor lock-in for UI components',
    ]
    for item in concerns:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 10. PERFORMANCE CONSIDERATIONS
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '10. Performance Considerations', level=1)

    add_styled_heading(doc, '10.1 Multi-Tenant Performance', level=2)

    doc.add_paragraph(
        'In a multi-tenant environment, performance isolation between tenants is essential to prevent '
        'one tenant\'s heavy usage from degrading service for others ("noisy neighbor" problem).'
    )

    perf_items = [
        ('Query Performance', 'The dual ORM approach (EF Core for CRUD, Dapper for stored procedures) provides good flexibility for performance optimization. Stored procedures can be tuned independently of application code.'),
        ('Connection Pooling', 'Standard ADO.NET connection pooling is used. At scale (100+ tenants), connection pool exhaustion may occur without pool partitioning or connection limits per tenant.'),
        ('Caching', 'No caching layer observed. Adding Redis or in-memory caching for frequently accessed tenant configurations and reference data would significantly improve response times.'),
        ('Blazor Server SignalR', 'Blazor Server maintains a persistent SignalR connection per user session. At 1,000+ concurrent users, this creates significant memory and connection pressure on the server.'),
        ('Blazor WebAssembly', 'The hybrid rendering model allows offloading interactive components to the client. This is a positive architectural choice for reducing server load.'),
    ]

    for title, desc in perf_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 11. RECOMMENDED AZURE CONFIGURATION
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '11. Recommended Azure Configuration', level=1)

    doc.add_paragraph(
        'The following Azure architecture is recommended for OrbaTech CRM targeting production '
        'deployment in the Canadian and US markets. The design emphasizes data residency compliance '
        '(Canada Central primary region), security best practices, and cost-effective scaling.'
    )

    doc.add_picture(chart_paths['azure_recommended'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run('Figure 5: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Recommended Azure architecture for OrbaTech CRM production deployment.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_styled_heading(doc, '11.1 Core Services', level=2)

    services = [
        ('Azure App Service (P1v3)', 'Primary compute for the Blazor application. P1v3 provides 2 vCPUs, 8 GB RAM, and supports deployment slots for zero-downtime updates. Auto-scale from 2 to 6 instances based on CPU/memory thresholds.'),
        ('Azure SQL Database', 'Start with Standard S2 (50 DTU) tier. Migrate to Elastic Pool as tenant count grows beyond 100 to share resources across tenant-specific workloads efficiently.'),
        ('Azure Front Door', 'Global load balancing with SSL termination, Web Application Firewall (WAF), and geographic routing between Canada Central and East US regions.'),
        ('Azure Blob Storage', 'Replace local filesystem for file attachments. Standard LRS tier with lifecycle management policies for cost optimization.'),
        ('Azure Key Vault', 'Centralized secrets management with managed identity authentication. Stores connection strings, API keys, and encryption keys.'),
        ('Azure Entra ID B2C', 'Cloud-native identity management replacing ASP.NET Core Identity. Supports SSO, MFA, and social login providers. Scales automatically with user growth.'),
        ('Azure Communication Services', 'Replace the Windows-based EmailService with a managed email delivery service. Supports email, SMS, and push notifications.'),
        ('Azure Monitor + Application Insights', 'Full-stack monitoring with distributed tracing, live metrics, and AI-powered anomaly detection. Log Analytics workspace for centralized log aggregation.'),
        ('Azure Redis Cache', 'Caching layer for tenant configuration, session state, and frequently accessed reference data. Basic C0 tier at startup, Standard C1 at growth.'),
    ]

    for title, desc in services:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    add_styled_heading(doc, '11.2 Security Configuration', level=2)

    sec_items = [
        'Managed Identities for all service-to-service authentication (no stored credentials)',
        'Private Endpoints for SQL Database, Key Vault, and Redis (no public network access)',
        'SQL Server Row-Level Security (RLS) policies for tenant isolation',
        'Azure RBAC with least-privilege access for all deployment identities',
        'Azure Policy for compliance enforcement (e.g., data residency, encryption)',
    ]
    for item in sec_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 12. AZURE COST PROJECTIONS
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '12. Azure Cost Projections', level=1)

    doc.add_paragraph(
        'The following cost estimates are based on Azure Canada Central region pricing as of April 2026. '
        'Costs are projected across three scaling tiers representing typical growth stages for an SMB CRM platform.'
    )

    doc.add_picture(chart_paths['azure_costs'], width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run('Figure 6: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Estimated monthly Azure costs across three scaling tiers. '
                     'Costs grow sub-linearly due to elastic resource sharing at scale.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Cost breakdown table
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Component', '10 Tenants', '100 Tenants', '1,000 Tenants']):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    cost_rows = [
        ('Compute (App Service)', '$50', '$300', '$1,800'),
        ('SQL Database', '$30', '$250', '$1,500'),
        ('Storage & CDN', '$5', '$30', '$200'),
        ('Identity (Entra ID B2C)', '$0', '$50', '$500'),
        ('Monitoring & Logging', '$10', '$50', '$200'),
        ('Networking & DNS', '$10', '$30', '$100'),
        ('Email Service', '$0', '$20', '$100'),
        ('Backup & DR', '$5', '$50', '$300'),
        ('TOTAL', '$110/mo', '$780/mo', '$4,700/mo'),
    ]
    for i, row in enumerate(cost_rows):
        for j, val in enumerate(row):
            table.cell(i+1, j).text = val
            for run in table.cell(i+1, j).paragraphs[0].runs:
                run.font.size = Pt(9)
        if i == len(cost_rows) - 1:  # Total row
            for j in range(4):
                for run in table.cell(i+1, j).paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: These estimates assume Azure Reserved Instances for compute (1-year commitment, ~35% savings) '
        'and standard support. Costs for DevExpress licensing (~$2,400/year per developer seat) are not included. '
        'Per-tenant cost at scale: approximately $4.70/tenant/month at 1,000 tenants.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 13. RECOMMENDED TECHNICAL PRIORITIES
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '13. Recommended Technical Priorities', level=1)

    doc.add_paragraph(
        'The following roadmap sequences the most impactful technical improvements, '
        'ordered by risk reduction and business value. Items are grouped into four priority levels.'
    )

    doc.add_picture(chart_paths['roadmap'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run('Figure 7: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Recommended technical priority roadmap showing parallel and sequential work streams over 12 months.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    priorities = [
        ('P0 — Critical (Months 1-3)', [
            ('Security Remediation', 'Remove all committed credentials from repository history (use git filter-repo). Rotate all exposed secrets. Implement Azure Key Vault for secrets management. Enable HTTPS redirection.'),
            ('CI/CD Pipeline', 'Implement GitHub Actions workflows for build, test, and deployment. Include vulnerability scanning (Dependabot). Add branch protection rules requiring PR reviews.'),
            ('Automated Testing', 'Establish unit test coverage for repository and service layers. Add integration tests for multi-tenant data isolation. Target minimum 60% code coverage.'),
        ]),
        ('P1 — High (Months 2-6)', [
            ('Multi-Tenant RLS', 'Implement SQL Server Row-Level Security policies. Add EF Core global query filters. Conduct cross-tenant penetration testing.'),
            ('Containerization', 'Create Dockerfile for the main application and EmailService. Implement docker-compose for local development. Prepare for Azure Container Apps or App Service deployment.'),
            ('Azure Deployment', 'Implement Infrastructure as Code (Bicep or Terraform). Set up staging and production environments. Configure deployment slots for zero-downtime releases.'),
            ('API Authentication', 'Implement OAuth 2.0 / OpenID Connect for API endpoints. Add rate limiting. Implement API versioning strategy.'),
        ]),
        ('P2 — Medium (Months 5-8)', [
            ('Localization (EN/FR)', 'Implement resource-based localization for all user-facing strings. Add French Canadian (fr-CA) translations. Plan for Spanish (es) as next language.'),
            ('Performance Optimization', 'Add Redis caching layer. Implement connection pool management per tenant. Optimize stored procedures for high-volume queries. Add response compression.'),
        ]),
        ('P3 — Future (Months 7-12)', [
            ('Mobile Responsive', 'Optimize Blazor components for mobile viewports. Consider Progressive Web App (PWA) capabilities.'),
            ('AI/Automation', 'Add predictive lead scoring, email templates, and workflow automation. Consider integration with AI services for natural language processing.'),
            ('Agent Red Integration', 'Implement bidirectional data flow with Agent Red AI for AI-powered customer service enrichment of CRM data (see Annex B).'),
        ]),
    ]

    for group_title, items in priorities:
        add_styled_heading(doc, group_title, level=2)
        for title, desc in items:
            p = doc.add_paragraph()
            run = p.add_run(f'{title}: ')
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(desc)
            run.font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # ANNEX A
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, 'Annex A: OrbaTech vs Agent Red Azure Deployment', level=1)

    doc.add_paragraph(
        'This annex compares the recommended Azure configuration for OrbaTech CRM against '
        'Remaker Digital\'s Agent Red platform deployment. The comparison illustrates different '
        'approaches to similar cloud architecture challenges and identifies synergy opportunities.'
    )

    doc.add_picture(chart_paths['annex_comparison'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run('Figure 8: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Side-by-side comparison of recommended OrbaTech and current Agent Red Azure deployments.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_styled_heading(doc, 'A.1 Key Architectural Differences', level=2)

    diff_items = [
        ('Compute Model', 'OrbaTech is recommended for Azure App Service (PaaS) for simplicity and managed infrastructure. Agent Red uses Azure Container Apps for microservice flexibility and independent scaling of 8+ service containers.'),
        ('Database Strategy', 'OrbaTech uses relational SQL Server, appropriate for structured CRM data with complex queries and reporting. Agent Red uses Cosmos DB for globally distributed, schema-flexible document storage optimized for real-time AI conversations.'),
        ('Identity Approach', 'OrbaTech should adopt Entra ID B2C for managed consumer identity. Agent Red implements custom identity with widget keys and API keys for embedded SaaS scenarios.'),
        ('Tenant Isolation', 'Both platforms use application-level tenant isolation. Agent Red additionally uses Cosmos DB partition keys for physical data separation. OrbaTech is recommended to add SQL RLS.'),
        ('CI/CD Maturity', 'Agent Red has a mature 13-phase pipeline with automated testing, security scanning, and staged deployment. OrbaTech needs to build this capability from the ground up.'),
    ]

    for title, desc in diff_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    add_styled_heading(doc, 'A.2 Shared Infrastructure Opportunities', level=2)

    shared_items = [
        'Shared Azure Key Vault patterns for secrets management',
        'Common Azure Monitor / Application Insights workspace for cross-platform observability',
        'Shared Azure Communication Services instance for email and SMS delivery',
        'Common CI/CD pipeline templates (GitHub Actions) adaptable to both .NET and Python stacks',
        'Shared Azure Front Door configuration for unified SSL and WAF management',
    ]
    for item in shared_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # ANNEX B
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, 'Annex B: OrbaTech + Agent Red Integration', level=1)

    doc.add_paragraph(
        'This annex describes a bidirectional integration architecture between OrbaTech CRM and '
        'Agent Red AI Customer Service. The integration enables AI-informed customer relationship '
        'management and CRM-enriched AI customer interactions.'
    )

    doc.add_picture(chart_paths['integration'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run('Figure 9: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Bidirectional integration architecture between OrbaTech CRM and Agent Red AI.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_styled_heading(doc, 'B.1 CRM Data to Agent Red (Left-to-Right)', level=2)

    ltr_items = [
        ('Contact Context', 'Customer profiles, purchase history, and communication preferences flow from OrbaTech to Agent Red, enabling personalized AI responses.'),
        ('Opportunity Status', 'Active deals and pipeline stages inform Agent Red\'s conversation routing, ensuring high-value prospects receive priority handling.'),
        ('Activity History', 'Recent emails, calls, and meetings provide Agent Red with conversation context, reducing customer frustration from repeating information.'),
        ('Custom Fields', 'Tenant-specific custom fields (e.g., industry, preferred language) enable Agent Red to adapt its responses to each business\'s unique CRM taxonomy.'),
    ]

    for title, desc in ltr_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    add_styled_heading(doc, 'B.2 Agent Red to CRM (Right-to-Left)', level=2)

    rtl_items = [
        ('Conversation Summaries', 'AI-generated summaries of customer interactions are written back to CRM contact records as activities, creating a complete interaction timeline.'),
        ('Sentiment Analysis', 'Customer sentiment scores from AI conversations are attached to opportunity records, providing sales teams with early warning on at-risk deals.'),
        ('Quality Metrics', 'Per-interaction quality scores help sales managers identify training opportunities and customer satisfaction trends.'),
        ('Lead Qualification', 'AI-assessed purchase intent signals from customer conversations create or update opportunity records in the CRM pipeline.'),
    ]

    for title, desc in rtl_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    add_styled_heading(doc, 'B.3 Integration Technical Requirements', level=2)

    tech_reqs = [
        'OAuth 2.0 service-to-service authentication with client credentials flow',
        'REST API endpoints on both platforms with OpenAPI/Swagger documentation',
        'Webhook-based event notifications for real-time data synchronization',
        'Tenant mapping table linking OrbaTech TenantId to Agent Red tenant identifiers',
        'Data transformation layer for field mapping between CRM entities and Agent Red conversation models',
        'Rate limiting and circuit breaker patterns for fault-tolerant integration',
        'Audit logging of all cross-platform data exchanges for compliance and debugging',
    ]
    for item in tech_reqs:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ── Closing ──
    doc.add_page_break()
    add_styled_heading(doc, 'Disclaimer', level=1)
    doc.add_paragraph(
        'This report was prepared by Remaker Digital based on publicly available information '
        'in the OrbaTech GitHub repository as of April 7, 2026. The findings and recommendations '
        'reflect the state of the codebase at that point in time. Technology assessments, cost '
        'estimates, and competitive comparisons are based on publicly available data and are provided '
        'for informational purposes. Azure pricing estimates are approximate and subject to change.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('© 2026 Remaker Digital, a DBA of VanDusen & Palmeter, LLC. All rights reserved.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Save ──
    print(f"Saving to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    print(f"Report saved: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == '__main__':
    build_report()
