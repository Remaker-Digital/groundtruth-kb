#!/usr/bin/env python3
"""
Agent Red Technical Evaluation Report Generator
Produces a professional DOCX with embedded charts and diagrams.
Mirror structure of OrbaTech report with reversed subject/annex roles.

© 2026 Remaker Digital, a DBA of VanDusen & Palmeter, LLC. All rights reserved.
"""

import os
import math

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.ticker import FuncFormatter
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────────
BASE = r"E:\Claude-Playground\CLAUDE-PROJECTS\Agent Red Customer Engagement"
LOGO_PATH = os.path.join(BASE, "branding", "logo", "PNG", "NEW-BLOCK-LOGO-HORIZONTAL-LIGHT.png")
OUTPUT_PATH = os.path.join(BASE, "AgentRed-Technical-Evaluation-Report.docx")
CHART_DIR = os.path.join(BASE, "scripts", "_report_charts_ar")
os.makedirs(CHART_DIR, exist_ok=True)

# ── Color Palette ──────────────────────────────────────────────────────────
REMAKER_RED = "#FF3621"
DARK_GRAY = "#2D2D2D"
MID_GRAY = "#666666"
ACCENT_BLUE = "#2196F3"
ACCENT_GREEN = "#4CAF50"
ACCENT_AMBER = "#FFC107"
ACCENT_RED = "#F44336"
ACCENT_PURPLE = "#9C27B0"

# ── Chart Generation ────────────────────────────────────────────────────────

def save_chart(fig, name, dpi=200):
    path = os.path.join(CHART_DIR, f"{name}.png")
    fig.savefig(path, dpi=dpi, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    return path


def chart_maturity_radar():
    """Radar chart: Agent Red maturity across 8 dimensions."""
    categories = [
        'Tech Stack', 'Architecture', 'Code Quality',
        'Testing', 'CI/CD', 'Security',
        'Documentation', 'Deployment'
    ]
    scores = [9, 9, 8, 8, 8, 9, 7, 9]  # out of 10

    N = len(categories)
    angles = [n / float(N) * 2 * math.pi for n in range(N)]
    angles += angles[:1]
    scores_plot = scores + scores[:1]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    ax.set_theta_offset(math.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=11, fontweight='bold', color=DARK_GRAY)
    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(['2', '4', '6', '8', '10'], fontsize=9, color=MID_GRAY)
    ax.yaxis.grid(True, color='#E0E0E0', linewidth=0.5)
    ax.xaxis.grid(True, color='#E0E0E0', linewidth=0.5)
    ax.plot(angles, scores_plot, 'o-', linewidth=2.5, color=REMAKER_RED, markersize=8)
    ax.fill(angles, scores_plot, alpha=0.15, color=REMAKER_RED)
    ax.set_title('Agent Red Engineering Maturity Assessment', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=30)
    return save_chart(fig, 'maturity_radar')


def chart_competitive_comparison():
    """Horizontal bar chart: AI customer service competitive comparison."""
    features = [
        'AI Response Quality', 'Multi-Agent Pipeline', 'Real-Time Streaming',
        'Knowledge Base', 'Conversation History', 'Multi-Language',
        'Multi-Tenant Isolation', 'Embeddable Widget', 'Admin Dashboard',
        'API/Integrations', 'Analytics/Quality', 'Security/Compliance'
    ]
    features.reverse()

    # Scores 0-5: 0=None, 1=Basic, 2=Functional, 3=Good, 4=Strong, 5=Enterprise
    data = {
        'Agent Red':       [4, 5, 5, 4, 4, 4, 5, 4, 4, 4, 4, 4],
        'Salesforce SC':   [4, 3, 4, 5, 5, 5, 5, 3, 5, 5, 5, 5],
        'Zendesk AI':      [4, 3, 4, 4, 5, 5, 5, 4, 5, 5, 4, 5],
        'Intercom Fin':    [4, 2, 4, 4, 4, 4, 4, 5, 4, 4, 4, 4],
        'Tidio':           [3, 1, 3, 3, 3, 3, 3, 4, 3, 3, 2, 2],
    }
    for k in data:
        data[k].reverse()

    fig, ax = plt.subplots(figsize=(12, 8))
    y_pos = np.arange(len(features))
    bar_height = 0.15
    colors = [REMAKER_RED, '#1A237E', '#00695C', '#FF6F00', '#6A1B9A']

    for i, (name, scores) in enumerate(data.items()):
        offset = (i - 2) * bar_height
        ax.barh(y_pos + offset, scores, bar_height, label=name,
                color=colors[i], alpha=0.85, edgecolor='white', linewidth=0.5)

    ax.set_xlabel('Capability Level', fontsize=11, color=DARK_GRAY)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(features, fontsize=10, color=DARK_GRAY)
    ax.set_xlim(0, 5.5)
    ax.set_xticks([0, 1, 2, 3, 4, 5])
    ax.set_xticklabels(['None', 'Basic', 'Functional', 'Good', 'Strong', 'Enterprise'], fontsize=9)
    ax.legend(loc='lower right', fontsize=10, framealpha=0.9)
    ax.set_title('AI Customer Service Feature Comparison', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='x', alpha=0.3)
    return save_chart(fig, 'competitive_comparison')


def chart_azure_costs():
    """Stacked bar chart: Agent Red Azure cost estimates at scale tiers."""
    tiers = ['10 Tenants\n(Startup)', '100 Tenants\n(Growth)', '1,000 Tenants\n(Scale)']

    costs = {
        'Container Apps (9 svcs)': [220, 900, 5500],
        'Cosmos DB':               [50, 400, 2800],
        'Azure OpenAI':            [100, 800, 6000],
        'Key Vault + Encryption':  [10, 30, 100],
        'NATS + Redis':            [30, 80, 300],
        'Monitoring (Langfuse+AI)': [20, 80, 350],
        'Blob Storage + CDN':      [10, 40, 200],
        'Comms (Email/SMS)':       [5, 50, 400],
    }

    fig, ax = plt.subplots(figsize=(10, 7))
    x = np.arange(len(tiers))
    width = 0.5
    bottom = np.zeros(len(tiers))
    colors = ['#1565C0', '#2196F3', '#FF6F00', '#4CAF50',
              '#B71C1C', '#E65100', '#64B5F6', '#00695C']

    for i, (label, vals) in enumerate(costs.items()):
        ax.bar(x, vals, width, label=label, bottom=bottom,
               color=colors[i], edgecolor='white', linewidth=0.5)
        bottom += np.array(vals)

    totals = [sum(v[i] for v in costs.values()) for i in range(3)]
    for i, total in enumerate(totals):
        ax.text(i, total + 50, f'${total:,}/mo', ha='center', va='bottom',
                fontsize=12, fontweight='bold', color=DARK_GRAY)

    ax.set_ylabel('Monthly Cost (USD)', fontsize=11, color=DARK_GRAY)
    ax.set_xticks(x)
    ax.set_xticklabels(tiers, fontsize=11, fontweight='bold', color=DARK_GRAY)
    ax.legend(loc='upper left', fontsize=9, framealpha=0.9)
    ax.set_title('Agent Red Estimated Azure Monthly Costs by Scale Tier', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.yaxis.set_major_formatter(FuncFormatter(lambda x, _: f'${x:,.0f}'))
    return save_chart(fig, 'azure_costs')


def chart_architecture_current():
    """Architecture diagram: Agent Red multi-agent pipeline."""
    fig, ax = plt.subplots(figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_title('Agent Red Multi-Agent Pipeline Architecture', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)

    def draw_box(x, y, w, h, label, color, sublabel=None, fontsize=10):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                                        facecolor=color, edgecolor='#333333',
                                        linewidth=1.5, alpha=0.85)
        ax.add_patch(rect)
        is_light = color in ['#FFF9C4', '#E3F2FD', '#F3E5F5', '#E8F5E9', '#FBE9E7']
        ax.text(x + w/2, y + h/2 + (0.12 if sublabel else 0), label,
                ha='center', va='center', fontsize=fontsize, fontweight='bold',
                color=DARK_GRAY if is_light else 'white')
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.2, sublabel, ha='center', va='center',
                    fontsize=8, color=MID_GRAY if is_light else '#DDDDDD')

    # Client tier
    draw_box(0.5, 8.5, 4, 1.0, 'Chat Widget', '#D84315', 'Preact | Shadow DOM | iframe', fontsize=12)
    draw_box(5, 8.5, 4, 1.0, 'Admin SPAs (x3)', '#1565C0', 'React | Mantine | Vite', fontsize=12)
    draw_box(9.5, 8.5, 4, 1.0, 'Shopify App', '#2E7D32', 'App Bridge | Embedded', fontsize=12)

    # Gateway
    draw_box(0.5, 7.0, 13, 1.0, 'API Gateway (FastAPI + Uvicorn)', REMAKER_RED,
             '52 Routers | SSE Streaming | JWT + API Key Auth | Rate Limiting', fontsize=12)

    # Agent pipeline
    draw_box(0.5, 5.2, 2.5, 1.3, 'Intent\nClassifier', '#6A1B9A', 'GPT-4o-mini')
    draw_box(3.5, 5.2, 2.5, 1.3, 'Knowledge\nRetrieval', '#1565C0', 'Embeddings + RAG')
    draw_box(6.5, 5.2, 2.5, 1.3, 'Response\nGenerator', '#2E7D32', 'GPT-4o')
    draw_box(9.5, 5.2, 2, 1.3, 'Critic\nSupervisor', '#E65100', 'Quality Gate')
    draw_box(12, 5.2, 1.5, 1.3, 'Escalation\nAgent', '#B71C1C')

    # Message bus
    draw_box(0.5, 3.8, 13, 0.9, 'NATS JetStream — Agent-to-Agent Message Bus', '#37474F',
             'Per-Tenant Isolation | Durable Subscriptions', fontsize=11)

    # Plugin agents
    draw_box(0.5, 2.3, 2, 1.0, 'Sales\nAgent', '#7B1FA2')
    draw_box(3, 2.3, 2, 1.0, 'Campaigns\nAgent', '#7B1FA2')
    draw_box(5.5, 2.3, 2, 1.0, 'Schedule\nAgent', '#7B1FA2')
    draw_box(8, 2.3, 2, 1.0, 'Gateway\nAgent', '#7B1FA2')
    draw_box(10.5, 2.3, 3, 1.0, '+14 Plugin Agents', '#7B1FA2', 'YAML Config-Driven')

    # Data tier
    draw_box(0.5, 0.5, 3.5, 1.3, 'Azure Cosmos DB', '#1A237E',
             '20+ Collections | Tenant Partitions')
    draw_box(4.5, 0.5, 3, 1.3, 'Azure Key Vault', '#2E7D32',
             'Envelope Encryption | HSM KEK')
    draw_box(8, 0.5, 2.5, 1.3, 'Redis', '#B71C1C', 'Rate Limiting | Cache')
    draw_box(11, 0.5, 2.5, 1.3, 'Azure Blob', '#455A64', 'Attachments | Exports')

    # Arrows
    arrow_props = dict(arrowstyle='->', color=MID_GRAY, lw=1.5)
    for x_off in [2.5, 7, 11.5]:
        ax.annotate('', xy=(x_off, 8.5), xytext=(x_off, 8.0), arrowprops=arrow_props)
    ax.annotate('', xy=(7, 7.0), xytext=(7, 6.5), arrowprops=arrow_props)
    for x_off in [1.75, 4.75, 7.75]:
        ax.annotate('', xy=(x_off, 5.2), xytext=(x_off, 4.7), arrowprops=arrow_props)

    return save_chart(fig, 'architecture_current')


def chart_azure_deployment():
    """Agent Red current Azure deployment diagram."""
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title('Agent Red Current Azure Deployment', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)

    def draw_box(x, y, w, h, label, color, sublabel=None, fontsize=10):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                                        facecolor=color, edgecolor='#333333',
                                        linewidth=1.5, alpha=0.85)
        ax.add_patch(rect)
        is_light = color in ['#FFF9C4', '#E3F2FD', '#F3E5F5', '#E8F5E9', '#FBE9E7']
        ax.text(x + w/2, y + h/2 + (0.12 if sublabel else 0), label,
                ha='center', va='center', fontsize=fontsize, fontweight='bold',
                color=DARK_GRAY if is_light else 'white')
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.2, sublabel, ha='center', va='center',
                    fontsize=8, color=MID_GRAY if is_light else '#DDDDDD')

    # Ingress
    draw_box(0.5, 7.5, 13, 1.0, 'Azure Container Apps Ingress + Custom Domains', '#0D47A1',
             'SSL Termination | CORS | Health Probes', fontsize=12)

    # Container Apps
    draw_box(0.5, 5.8, 3, 1.2, 'API Gateway', REMAKER_RED, '0.5 CPU | 2-8 replicas')
    draw_box(4, 5.8, 3, 1.2, 'Intent Classifier', '#6A1B9A', '0.5 CPU | 2-6 replicas')
    draw_box(7.5, 5.8, 3, 1.2, 'Response Generator', '#2E7D32', '1.0 CPU | 2-10 replicas')
    draw_box(11, 5.8, 2.5, 1.2, 'Critic + 4 more', '#455A64', 'Scaled independently')

    # Services
    draw_box(0.5, 4.0, 3, 1.2, 'Cosmos DB', '#1A237E', 'Serverless + Provisioned\n20+ Collections')
    draw_box(4, 4.0, 3, 1.2, 'Azure OpenAI', '#FF6F00', 'GPT-4o | GPT-4o-mini\nEmbeddings')
    draw_box(7.5, 4.0, 3, 1.2, 'Azure Key Vault', '#2E7D32', 'HSM KEK | DEKs\nManaged Identity')
    draw_box(11, 4.0, 2.5, 1.2, 'NATS JetStream', '#37474F', 'Agent Bus\n2 replicas')

    # Supporting
    draw_box(0.5, 2.3, 3, 1.2, 'Azure Monitor', '#E65100', 'App Insights + Langfuse\nLog Analytics')
    draw_box(4, 2.3, 3, 1.2, 'Redis', '#B71C1C', 'Rate Limiting\nTenant Config Cache')
    draw_box(7.5, 2.3, 3, 1.2, 'Azure Comm Svc', '#00695C', 'Email + SMS OTP')
    draw_box(11, 2.3, 2.5, 1.2, 'Azure Blob', '#455A64', 'Exports + Backups')

    # IaC + CI/CD
    draw_box(0.5, 0.8, 6, 1.0, 'Terraform IaC (7 .tf files) | GitHub Actions (8 workflows)',
             '#E8F5E9', fontsize=10)
    draw_box(7, 0.8, 6.5, 1.0, 'ACR Registry | 3 Dockerfiles | Managed Identity RBAC',
             '#FBE9E7', fontsize=10)

    return save_chart(fig, 'azure_deployment')


def chart_priority_roadmap():
    """Gantt-style chart for Agent Red recommended priorities."""
    fig, ax = plt.subplots(figsize=(12, 7))

    tasks = [
        ('Production Monitoring Enhancement', 0, 2, ACCENT_AMBER, 'P1 — High'),
        ('Load/Stress Testing at Scale', 1, 3, ACCENT_AMBER, 'P1 — High'),
        ('SOC 2 Type II Preparation', 1, 6, ACCENT_AMBER, 'P1 — High'),
        ('Public API Documentation', 2, 4, ACCENT_AMBER, 'P1 — High'),
        ('Multi-Region DR', 3, 6, ACCENT_GREEN, 'P2 — Medium'),
        ('French/Spanish Localization QA', 2, 4, ACCENT_GREEN, 'P2 — Medium'),
        ('Mobile-Responsive Admin SPA', 4, 7, ACCENT_GREEN, 'P2 — Medium'),
        ('Partner Integration SDK', 5, 8, ACCENT_PURPLE, 'P3 — Future'),
        ('Self-Serve Onboarding', 6, 9, ACCENT_PURPLE, 'P3 — Future'),
        ('Marketplace / App Store', 8, 12, ACCENT_PURPLE, 'P3 — Future'),
        ('OrbaTech CRM Integration', 4, 7, ACCENT_BLUE, 'Partnership'),
    ]
    tasks.reverse()

    for i, (name, start, end, color, priority) in enumerate(tasks):
        ax.barh(i, end - start, left=start, height=0.6, color=color,
                alpha=0.8, edgecolor='white', linewidth=0.5)
        ax.text(end + 0.2, i, priority, va='center', fontsize=8, color=MID_GRAY)

    ax.set_yticks(range(len(tasks)))
    ax.set_yticklabels([t[0] for t in tasks], fontsize=10, color=DARK_GRAY)
    ax.set_xlabel('Months', fontsize=11, color=DARK_GRAY)
    ax.set_title('Agent Red Recommended Technical Priority Roadmap', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_xlim(0, 13)
    ax.grid(axis='x', alpha=0.3)

    legend_elements = [
        mpatches.Patch(facecolor=ACCENT_AMBER, label='P1 — High (Months 1-6)'),
        mpatches.Patch(facecolor=ACCENT_GREEN, label='P2 — Medium (Months 2-7)'),
        mpatches.Patch(facecolor=ACCENT_PURPLE, label='P3 — Future (Months 5-12)'),
        mpatches.Patch(facecolor=ACCENT_BLUE, label='Partnership Opportunity'),
    ]
    ax.legend(handles=legend_elements, loc='lower right', fontsize=9)
    return save_chart(fig, 'priority_roadmap')


def chart_security_heatmap():
    """Heatmap for Agent Red security posture."""
    categories = [
        'Authentication', 'Authorization', 'Data Encryption',
        'Secrets Mgmt', 'Input Validation', 'API Security',
        'Tenant Isolation', 'Audit Logging', 'HTTPS/TLS',
        'Dependency Scanning'
    ]
    scores = [5, 4, 5, 5, 4, 4, 5, 4, 5, 4]

    fig, ax = plt.subplots(figsize=(10, 5))
    colors_map = {1: '#D32F2F', 2: '#F57C00', 3: '#FBC02D', 4: '#66BB6A', 5: '#2E7D32'}
    labels_map = {1: 'Critical Gap', 2: 'Needs Work', 3: 'Partial', 4: 'Adequate', 5: 'Strong'}

    ax.barh(range(len(categories)), scores, color=[colors_map[s] for s in scores],
            height=0.6, edgecolor='white', linewidth=1)

    for i, (score, cat) in enumerate(zip(scores, categories)):
        ax.text(score + 0.1, i, labels_map[score], va='center', fontsize=9, color=MID_GRAY)

    ax.set_yticks(range(len(categories)))
    ax.set_yticklabels(categories, fontsize=10, color=DARK_GRAY)
    ax.set_xlim(0, 6)
    ax.set_xticks([1, 2, 3, 4, 5])
    ax.set_xticklabels(['Critical\nGap', 'Needs\nWork', 'Partial', 'Adequate', 'Strong'], fontsize=9)
    ax.set_title('Security Posture Assessment', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='x', alpha=0.2)
    return save_chart(fig, 'security_heatmap')


def chart_integration_architecture():
    """Bidirectional integration: Agent Red <-> OrbaTech CRM."""
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_title('Bidirectional Integration: Agent Red AI ↔ OrbaTech CRM', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)

    def draw_box(x, y, w, h, label, color, sublabel=None, fontsize=10):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                                        facecolor=color, edgecolor='#333333',
                                        linewidth=1.5, alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.12 if sublabel else 0), label,
                ha='center', va='center', fontsize=fontsize, fontweight='bold', color='white')
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.3, sublabel, ha='center', va='center',
                    fontsize=8, color='#DDDDDD')

    # Agent Red side (LEFT — subject of this report)
    draw_box(0.3, 5.0, 5, 1.3, 'Agent Red AI', REMAKER_RED,
             'Python | Azure | Cosmos DB', fontsize=13)
    draw_box(0.3, 3.2, 2.3, 1.3, 'AI Customer\nService', '#D84315')
    draw_box(3, 3.2, 2.3, 1.3, 'Knowledge\nBase', '#D84315')
    draw_box(0.3, 1.5, 2.3, 1.3, 'Conversation\nHistory', '#D84315')
    draw_box(3, 1.5, 2.3, 1.3, 'Quality\nAnalytics', '#D84315')

    # OrbaTech side (RIGHT — annex subject)
    draw_box(7.7, 5.0, 5, 1.3, 'OrbaTech CRM', '#1565C0',
             '.NET 8 | Blazor | SQL Server', fontsize=13)
    draw_box(7.7, 3.2, 2.3, 1.3, 'Contacts &\nAccounts', '#1976D2')
    draw_box(10.4, 3.2, 2.3, 1.3, 'Opportunities\n& Pipeline', '#1976D2')
    draw_box(7.7, 1.5, 2.3, 1.3, 'Activities &\nEmails', '#1976D2')
    draw_box(10.4, 1.5, 2.3, 1.3, 'Custom\nFields', '#1976D2')

    # Integration hub
    draw_box(5.6, 3.5, 1.8, 2.5, 'Integration\nHub', '#2E7D32',
             'REST API\nWebhooks\nOAuth 2.0', fontsize=11)

    arrow_r = dict(arrowstyle='->', color=REMAKER_RED, lw=2.5)
    arrow_l = dict(arrowstyle='->', color='#1565C0', lw=2.5)
    ax.annotate('', xy=(5.6, 5.2), xytext=(5.3, 5.2), arrowprops=arrow_r)
    ax.annotate('', xy=(7.7, 5.2), xytext=(7.4, 5.2), arrowprops=arrow_r)
    ax.annotate('', xy=(7.4, 4.6), xytext=(7.7, 4.6), arrowprops=arrow_l)
    ax.annotate('', xy=(5.3, 4.6), xytext=(5.6, 4.6), arrowprops=arrow_l)

    ax.text(6.5, 6.5, 'AI Insights → CRM Records', fontsize=10, ha='center',
            color=REMAKER_RED, fontweight='bold')
    ax.text(6.5, 0.7, 'CRM Data → AI Context', fontsize=10, ha='center',
            color='#1565C0', fontweight='bold')
    return save_chart(fig, 'integration_architecture')


def chart_annex_comparison():
    """Side-by-side comparison table as visual chart."""
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.axis('off')
    ax.set_title('Annex A: Agent Red Azure Deployment vs OrbaTech Recommended Configuration',
                 fontsize=13, fontweight='bold', color=DARK_GRAY, pad=15)

    columns = ['Component', 'Agent Red\n(Current)', 'OrbaTech\n(Recommended)']
    rows = [
        ['Compute', 'Container Apps\n9 containers, auto-scale', 'App Service P1v3\nAuto-scale 2-6'],
        ['Database', 'Cosmos DB\nServerless + Provisioned', 'Azure SQL\nElastic Pool (DTU)'],
        ['Identity', 'Custom Identity\nWidget Keys + API Keys + JWT', 'Entra ID B2C\nCookie + OAuth'],
        ['File Storage', 'Azure Blob\nStandard LRS', 'Azure Blob\nStandard LRS'],
        ['Secrets', 'Azure Key Vault\nManaged Identity + KEK/DEK', 'Azure Key Vault\nManaged Identity'],
        ['CDN/Gateway', 'Container Apps Ingress\nCustom domains', 'Azure Front Door\nWAF + Geo-routing'],
        ['Email/SMS', 'Azure Comm Services\nEmail + SMS OTP', 'Azure Comm Services\nEmail only'],
        ['Monitoring', 'App Insights + Langfuse\nCustom Observability', 'App Insights +\nLog Analytics'],
        ['CI/CD', 'GitHub Actions\n8 workflows, 5-shard tests', 'GitHub Actions\n(To be built)'],
        ['Multi-Tenant', 'Cosmos Partition Keys\n+ App-Level + Encryption', 'App-Level +\nSQL RLS (planned)'],
    ]

    table = ax.table(cellText=rows, colLabels=columns,
                     cellLoc='center', loc='center', colWidths=[0.25, 0.35, 0.35])
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 2.2)

    for j in range(3):
        cell = table[0, j]
        cell.set_facecolor(REMAKER_RED if j == 1 else '#1565C0')
        cell.set_text_props(color='white', fontweight='bold', fontsize=10)
        cell.set_edgecolor('white')

    for i in range(1, len(rows) + 1):
        for j in range(3):
            cell = table[i, j]
            cell.set_facecolor('#F5F5F5' if i % 2 == 0 else 'white')
            cell.set_edgecolor('#E0E0E0')
            if j == 0:
                cell.set_text_props(fontweight='bold')

    return save_chart(fig, 'annex_comparison')


def chart_test_coverage():
    """Test distribution by type."""
    categories = ['Unit', 'Integration', 'E2E/\nBrowser', 'Contract/\nAPI', 'Security', 'Performance/\nLoad', 'Property/\nFuzzing', 'Visual/\nA11y']
    counts = [950, 3700, 120, 310, 80, 50, 200, 30]

    fig, ax = plt.subplots(figsize=(10, 5))
    colors = ['#1565C0', '#2196F3', '#4CAF50', '#FF6F00', '#D32F2F', '#9C27B0', '#00695C', '#795548']
    bars = ax.bar(range(len(categories)), counts, color=colors, edgecolor='white', linewidth=1, alpha=0.85)

    for bar, count in zip(bars, counts):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 30,
                f'{count:,}', ha='center', va='bottom', fontsize=10, fontweight='bold', color=DARK_GRAY)

    ax.set_xticks(range(len(categories)))
    ax.set_xticklabels(categories, fontsize=9, color=DARK_GRAY)
    ax.set_ylabel('Number of Tests', fontsize=11, color=DARK_GRAY)
    ax.set_title(f'Test Suite Distribution (Total: {sum(counts):,} tests)', fontsize=14,
                 fontweight='bold', color=DARK_GRAY, pad=15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='y', alpha=0.2)
    return save_chart(fig, 'test_coverage')


# ── DOCX Helpers ────────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading_elm)


def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return heading


def add_key_finding_box(doc, title, text, severity='info'):
    colors = {
        'critical': 'FF3621', 'high': 'F57C00', 'medium': 'FBC02D',
        'info': '2196F3', 'positive': '4CAF50',
    }
    color = colors.get(severity, '2196F3')
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F5F5F5')
    p_title = cell.paragraphs[0]
    run = p_title.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)
    p_text = cell.add_paragraph(text)
    p_text.style = doc.styles['Normal']
    for run in p_text.runs:
        run.font.size = Pt(10)
    doc.add_paragraph()


def add_figure_caption(doc, fig_num, text):
    p = doc.add_paragraph()
    run = p.add_run(f'Figure {fig_num}: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_detail_item(doc, title, desc):
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(desc)
    run.font.size = Pt(10)


# ── DOCX Report ────────────────────────────────────────────────────────────

def build_report():
    print("Generating charts...")
    charts = {
        'maturity': chart_maturity_radar(),
        'competitive': chart_competitive_comparison(),
        'azure_costs': chart_azure_costs(),
        'architecture': chart_architecture_current(),
        'azure_deploy': chart_azure_deployment(),
        'roadmap': chart_priority_roadmap(),
        'security': chart_security_heatmap(),
        'integration': chart_integration_architecture(),
        'annex_comparison': chart_annex_comparison(),
        'test_coverage': chart_test_coverage(),
    }
    print("Charts generated.")

    doc = Document()

    # ── Page setup + styles ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Agent Red')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0xFF, 0x36, 0x21)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Technical Evaluation Report')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Prepared by Remaker Digital')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('April 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = classification.add_run('CONFIDENTIAL')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xFF, 0x36, 0x21)
    run.bold = True

    doc.add_page_break()

    # ── Footer ──
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer_para.add_run()
        run.add_picture(LOGO_PATH, height=Inches(0.35))

    # ── Table of Contents ──
    add_styled_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '3. Technology Assessment',
        '4. Engineering Maturity',
        '5. Security Assessment',
        '6. Multi-Tenant Architecture',
        '7. Competitive Landscape',
        '8. Reliability & Serviceability',
        '9. Maintainability',
        '10. Performance Considerations',
        '11. Azure Deployment Configuration',
        '12. Azure Cost Projections',
        '13. Recommended Technical Priorities',
        'Annex A: Agent Red vs OrbaTech Recommended Azure Configuration',
        'Annex B: Agent Red + OrbaTech CRM Integration',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '1. Executive Summary', level=1)

    doc.add_paragraph(
        'This report provides an independent technical evaluation of the Agent Red AI Customer '
        'Service platform, a multi-agent AI system built on Python/FastAPI with Azure infrastructure. '
        'The evaluation was conducted by Remaker Digital as part of a partnership assessment with OrbaTech CRM.'
    )
    doc.add_paragraph(
        'Agent Red is an actively developed product (approximately 6 months in development, 711 commits) '
        'built by a single developer with extensive AI-assisted development (87.6% of commits co-authored '
        'with Claude). The platform provides AI-powered customer service for Shopify merchants and '
        'standalone businesses, featuring a multi-agent pipeline, admin dashboards, and an embeddable chat widget.'
    )

    add_key_finding_box(doc,
        'Overall Assessment',
        'Agent Red demonstrates strong engineering maturity across nearly all dimensions. The multi-agent '
        'pipeline architecture, comprehensive test suite (5,400+ tests), production-grade Azure deployment '
        '(9 containers, Terraform IaC), and defense-in-depth security posture (envelope encryption, '
        'managed identities, tenant-partitioned data) position the platform well for commercial production. '
        'Key areas for investment include SOC 2 compliance preparation, load testing at scale, and '
        'partner integration SDK development.',
        'positive'
    )

    doc.add_picture(charts['maturity'], width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 1, 'Engineering maturity across eight dimensions (scale of 1-10). '
                       'Agent Red scores 7+ across all dimensions, with particular strength in '
                       'Tech Stack, Architecture, Security, and Deployment.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 2. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '2. Project Overview', level=1)

    table = doc.add_table(rows=11, cols=2)
    table.style = 'Light Shading Accent 1'
    facts = [
        ('Repository', 'github.com/Remaker-Digital/agent-red-customer-engagement'),
        ('Description', 'AI-powered customer service SaaS — multi-agent pipeline, admin dashboards, AGNTCY platform'),
        ('Primary Language', 'Python 3.12 (backend), TypeScript (frontend)'),
        ('Project Age', '~6 months (since October 2025)'),
        ('Total Commits', '711'),
        ('Contributors', '1 developer + AI co-author (Claude Opus)'),
        ('Total Files', '2,213 (2,052 source)'),
        ('Test Files', '605 files, 5,400+ tests'),
        ('License', 'Proprietary (all rights reserved)'),
        ('Target Markets', 'Shopify merchants + standalone businesses (US/Canada)'),
        ('Website', 'Remaker Digital'),
    ]
    for i, (label, value) in enumerate(facts):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        for run in table.cell(i, 0).paragraphs[0].runs:
            run.bold = True
        for cell in table.row_cells(i):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph(
        'The project demonstrates exceptional development velocity with 711 commits over 6 months, '
        'averaging nearly 4 commits per day. The AI-assisted development model (87.6% co-authored commits) '
        'is a notable approach that has enabled a single developer to build and maintain a system of this '
        'complexity. The codebase comprises approximately 18 MB of source code across Python, TypeScript, '
        'HTML, CSS, T-SQL, and HCL (Terraform).'
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 3. TECHNOLOGY ASSESSMENT
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '3. Technology Assessment', level=1)

    add_key_finding_box(doc,
        'Technology Choices: Production-Grade',
        'The selection of Python/FastAPI for the backend, React/Preact for frontends, Azure Cosmos DB '
        'for multi-tenant data, and NATS JetStream for agent-to-agent communication represents a '
        'well-considered, cloud-native technology stack optimized for AI workloads and horizontal scaling.',
        'positive'
    )

    doc.add_picture(charts['architecture'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 2, 'Agent Red multi-agent pipeline architecture showing the 6 core agents, '
                       'plugin system (20 agents total), NATS message bus, and data tier.')
    doc.add_paragraph()

    add_styled_heading(doc, '3.1 Technology Stack Details', level=2)

    table = doc.add_table(rows=10, cols=3)
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Layer', 'Technology', 'Assessment']):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    stack_rows = [
        ('API Framework', 'FastAPI + Uvicorn\n(4 workers, tini PID 1)', 'High-performance async Python framework. Automatic OpenAPI docs. Pydantic v2 validation. Production-ready with tini signal handling.'),
        ('AI/ML', 'Azure OpenAI\n(GPT-4o, GPT-4o-mini)', 'Enterprise-grade AI models with Azure-managed scaling. Langfuse observability for prompt monitoring and cost tracking.'),
        ('Database', 'Azure Cosmos DB\n(20+ collections)', 'Globally distributed NoSQL with native partition-key tenant isolation. Serverless + provisioned throughput options. Continuous 7-day backup.'),
        ('Message Bus', 'NATS JetStream', 'High-performance agent-to-agent communication with per-tenant isolation, durable subscriptions, and at-least-once delivery.'),
        ('Frontend (Admin)', 'React 18.3 + Mantine\n(3 SPAs)', 'Three tailored admin interfaces: Shopify embedded, standalone, and platform operator. Shared component library.'),
        ('Frontend (Widget)', 'Preact 10.25\n(Shadow DOM + iframe)', 'Lightweight embeddable widget with CSS isolation (Shadow DOM launcher) and full DOM isolation (iframe panel). 8-language support.'),
        ('Encryption', 'AES-256-GCM\nAzure Key Vault HSM KEK', 'Per-tenant envelope encryption with HSM-backed master key. Auto-rotation. No plaintext sensitive data at rest.'),
        ('Infrastructure', 'Terraform + Docker\nAzure Container Apps', '9 independently scaled containers. Terraform IaC for environment management. 3 Dockerfiles with non-root user.'),
        ('CI/CD', 'GitHub Actions\n(8 workflows)', 'Build, test (5-shard parallel), lint, deploy pipelines. Dependency vulnerability scanning. OpenAPI compatibility checks.'),
    ]
    for i, (layer, tech, assessment) in enumerate(stack_rows):
        table.cell(i+1, 0).text = layer
        table.cell(i+1, 1).text = tech
        table.cell(i+1, 2).text = assessment
        for cell in table.row_cells(i+1):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 4. ENGINEERING MATURITY
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '4. Engineering Maturity', level=1)

    doc.add_paragraph(
        'Engineering maturity was assessed across eight dimensions using a 1-10 scale. '
        'Agent Red scores consistently high, reflecting mature engineering practices despite '
        'the relatively young project age.'
    )

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Dimension', 'Score', 'Key Observations']):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    maturity_rows = [
        ('Tech Stack', '9/10', 'Python 3.12, FastAPI, Cosmos DB, NATS, React, Preact. All current versions. Cloud-native choices.'),
        ('Architecture', '9/10', 'Multi-agent pipeline with mixin orchestrator. Plugin system. 3 SPAs + widget. Repository pattern with tenant-scoped base class.'),
        ('Code Quality', '8/10', 'Ruff linting (CI-enforced), import cycle detection, dead code analysis, cyclomatic complexity monitoring. ESLint for TypeScript.'),
        ('Testing', '8/10', '5,400+ tests across 10 test types. 5-shard parallel CI. Coverage gate. API fuzzing, mutation testing, property-based testing.'),
        ('CI/CD', '8/10', '8 GitHub Actions workflows. Automated build/test/deploy pipeline. Dependency vulnerability scanning. OpenAPI compatibility.'),
        ('Security', '9/10', 'Envelope encryption, HSM-backed keys, managed identities, input sanitization, pre-auth rate limiting, audit logging, PII tokenization.'),
        ('Documentation', '7/10', 'Extensive internal docs (CLAUDE.md system, knowledge DB with 2,092 specs). Admin guide site. Legal drafts. Could improve public API docs.'),
        ('Deployment', '9/10', '9 Azure Container Apps with Terraform IaC. 3 Dockerfiles. Staging + production. Managed identity RBAC. Health probes.'),
    ]
    for i, (dim, score, obs) in enumerate(maturity_rows):
        table.cell(i+1, 0).text = dim
        table.cell(i+1, 1).text = score
        table.cell(i+1, 2).text = obs
        for cell in table.row_cells(i+1):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    add_styled_heading(doc, '4.1 Test Suite', level=2)
    doc.add_paragraph(
        'Agent Red has an exceptionally comprehensive test suite covering 10 distinct testing methodologies. '
        'Tests run in 5 parallel shards in CI with a 75% coverage gate (ramping to 80%).'
    )

    doc.add_picture(charts['test_coverage'], width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 3, 'Test suite distribution by type. Integration tests form the largest category, '
                       'with significant investment in API fuzzing, property-based testing, and security testing.')

    add_styled_heading(doc, '4.2 Coding Standards', level=2)
    standards = [
        'Ruff linter enforced in CI (pycodestyle, pyflakes, isort, bugbear, simplify rules)',
        'Import cycle detection via custom AST walker (blocking in CI)',
        'pip-audit dependency vulnerability scanning (blocking in CI)',
        'Radon cyclomatic complexity monitoring (advisory)',
        'Vulture dead code detection (advisory)',
        'ESLint with react-hooks and jsx-a11y plugins for TypeScript',
        'Vale prose linting for documentation with custom AgentRed style rules',
        'Conventional commit message format (feat, fix, chore, bump, test, docs)',
    ]
    for item in standards:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 5. SECURITY ASSESSMENT
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '5. Security Assessment', level=1)

    add_key_finding_box(doc,
        'Security Posture: Strong',
        'Agent Red implements defense-in-depth security with per-tenant envelope encryption, '
        'Azure Key Vault HSM-backed keys, managed identities, input/output sanitization, '
        'pre-auth rate limiting, audit logging, and PII tokenization. No committed credentials '
        'or hardcoded secrets were found. The .gitignore properly excludes all sensitive files.',
        'positive'
    )

    doc.add_picture(charts['security'], width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 4, 'Security posture assessment across ten domains. All domains rate Adequate or Strong.')
    doc.add_paragraph()

    add_styled_heading(doc, '5.1 Security Implementation Details', level=2)

    sec_details = [
        ('Authentication (Strong)', 'Triple authentication model: Shopify JWT session tokens, SHA-256 hashed API keys via Key Vault, and publishable widget keys with pk_live_ prefix scoping. Additional support for MFA/TOTP, magic link auth, and Argon2 password hashing.'),
        ('Authorization (Adequate)', 'Per-tenant RBAC with TeamMemberRole enum. Platform admin separated via PLATFORM_ADMIN_TENANT_ID sentinel. Dedicated admin console. Area for improvement: more granular permission model.'),
        ('Data Encryption (Strong)', 'AES-256-GCM envelope encryption with Azure Key Vault HSM-backed master KEK. Per-tenant Data Encryption Keys (DEKs). Auto-rotation via Key Vault. All sensitive data encrypted at rest.'),
        ('Secrets Management (Strong)', 'All secrets in Azure Key Vault with managed identity access. No hardcoded keys policy enforced by guardrail regex (SPEC-1845). API key rotation endpoint. CMK auto-rotation.'),
        ('Tenant Isolation (Strong)', 'Cosmos DB partition keys enforce physical data separation at the database level. TenantScopedRepository base class structurally prevents cross-tenant access. Per-tenant NATS isolation for message routing.'),
        ('Audit Logging (Adequate)', 'Append-only audit log with 12 event types and 1-year retention. GDPR compliance with PII scrubbing, data export, cascading deletion, and consent management.'),
    ]

    for title, desc in sec_details:
        add_detail_item(doc, title, desc)

    add_styled_heading(doc, '5.2 Areas for Improvement', level=2)
    sec_improvements = [
        'SOC 2 Type II certification preparation (controls documentation, evidence collection)',
        'Penetration testing by an independent third party',
        'Web Application Firewall (WAF) at the ingress layer (currently using Container Apps native ingress)',
        'More granular RBAC beyond team member roles (field-level permissions)',
    ]
    for item in sec_improvements:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 6. MULTI-TENANT ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '6. Multi-Tenant Architecture', level=1)

    doc.add_paragraph(
        'Agent Red implements a deeply integrated multi-tenant architecture with isolation enforced at '
        'multiple layers: database partitioning, application-level repository pattern, message bus isolation, '
        'and per-tenant encryption.'
    )

    add_key_finding_box(doc,
        'Tenant Isolation: Enterprise-Grade',
        'The combination of Cosmos DB partition keys, TenantScopedRepository base class, per-tenant '
        'NATS isolation, and per-tenant envelope encryption provides defense-in-depth tenant isolation '
        'that exceeds typical SaaS security patterns.',
        'positive'
    )

    add_styled_heading(doc, '6.1 Isolation Layers', level=2)

    isolation_layers = [
        ('Database Layer', 'Cosmos DB partition keys use tenant_id as the partition key on all 20+ collections. Cross-partition queries are structurally impossible without explicit tenant_id specification.'),
        ('Application Layer', 'TenantScopedRepository base class enforces mandatory tenant_id on every database operation. All repositories inherit from this base class, making tenant isolation a structural guarantee rather than a convention.'),
        ('Message Bus Layer', 'NATS JetStream uses per-tenant subject namespacing for agent-to-agent communication, preventing message leakage between tenants.'),
        ('Encryption Layer', 'Per-tenant Data Encryption Keys (DEKs) ensure that even with database-level access, one tenant\'s data cannot be decrypted with another tenant\'s key.'),
        ('Observability Layer', 'Per-tenant OpenTelemetry correlation (tenant_id as span attribute) enables isolated monitoring and debugging without cross-tenant data exposure.'),
    ]

    for title, desc in isolation_layers:
        add_detail_item(doc, title, desc)

    add_styled_heading(doc, '6.2 Tenant Lifecycle', level=2)

    lifecycle_items = [
        '4 subscription tiers: trial, starter, professional, enterprise',
        'Lifecycle states: provisioning, active, past_due, grace_period (plus suspended/deprovisioned)',
        'Stripe-integrated billing with metered usage tracking',
        'Per-tenant rate limiting: 300 RPM ceiling, 10 RPM floor, platform admin exempt',
        'Domain index collection for tenant lookup without cross-partition queries',
        'License management per subscription tier',
    ]
    for item in lifecycle_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 7. COMPETITIVE LANDSCAPE
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '7. Competitive Landscape', level=1)

    doc.add_paragraph(
        'Agent Red operates in the AI-powered customer service market, competing against '
        'established players ranging from enterprise platforms to SMB-focused solutions. '
        'The following comparison positions Agent Red against key competitors.'
    )

    doc.add_picture(charts['competitive'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 5, 'Feature capability comparison across twelve AI customer service dimensions. '
                       'Agent Red shows strength in multi-agent pipeline, real-time streaming, and tenant isolation.')
    doc.add_paragraph()

    add_styled_heading(doc, '7.1 Competitor Profiles', level=2)

    competitors = [
        ('Salesforce Service Cloud', '$25-$500/user/mo', 'Market leader with Einstein AI. Comprehensive but complex and expensive. Strong enterprise adoption. AI features are add-ons to existing CRM.'),
        ('Zendesk AI', '$55-$169/agent/mo', 'Established support platform with AI agent capabilities. Strong ticket management and knowledge base. AI features expanding rapidly.'),
        ('Intercom Fin', '$29-$132/seat/mo', 'Leading conversational support platform. Fin AI agent handles common queries. Strong in-app messaging and product tours. Growing AI capabilities.'),
        ('Tidio', '$0-$59/mo', 'Budget-friendly chatbot platform for small businesses. Limited AI sophistication. Strong Shopify integration. Growing market share in micro-business segment.'),
    ]

    table = doc.add_table(rows=len(competitors)+1, cols=3)
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Competitor', 'Pricing', 'Market Position']):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    for i, (name, price, position) in enumerate(competitors):
        table.cell(i+1, 0).text = name
        table.cell(i+1, 1).text = price
        table.cell(i+1, 2).text = position
        for cell in table.row_cells(i+1):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    add_styled_heading(doc, '7.2 Agent Red Differentiation', level=2)

    diff_items = [
        'Multi-agent pipeline with specialized AI agents (intent, knowledge, response, critic) — deeper AI architecture than single-model competitors',
        'Critic-retraction pattern ensures AI response quality before delivery — unique quality gate not found in competitor products',
        'Per-tenant envelope encryption with HSM-backed keys — exceeds typical SaaS security practices',
        'Shopify-native integration via App Bridge — optimized for e-commerce merchants',
        'Embeddable widget with Shadow DOM isolation — minimal impact on merchant site performance',
        'AI-assisted development model (Claude co-authoring) — enables rapid feature development at scale',
        'CRM integration potential (OrbaTech partnership) — AI insights flowing into sales pipeline',
    ]
    for item in diff_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 8. RELIABILITY & SERVICEABILITY
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '8. Reliability & Serviceability', level=1)

    add_styled_heading(doc, '8.1 Reliability', level=2)

    rel_items = [
        ('Container health probes', 'All 9 containers have health check endpoints. Container Apps automatically restarts unhealthy instances and routes traffic away during failures.'),
        ('Auto-scaling', 'HTTP-triggered scaling (2-8 replicas for gateway, 2-10 for response generator) and NATS-triggered scaling for agent containers ensures capacity matches demand.'),
        ('Database resilience', 'Cosmos DB continuous backup with 7-day point-in-time restore (PITR). Serverless + provisioned throughput options prevent throttling under load.'),
        ('Message durability', 'NATS JetStream durable subscriptions provide at-least-once delivery guarantees for agent-to-agent communication.'),
        ('Critic-retraction', 'The critic agent can retract an AI response mid-stream if quality thresholds are not met, preventing low-quality responses from reaching customers.'),
        ('Areas for improvement', 'Multi-region disaster recovery is not yet implemented. Load testing at 1,000+ tenant scale has not been validated. Circuit breaker patterns could be more comprehensive.'),
    ]
    for title, desc in rel_items:
        add_detail_item(doc, title, desc)

    add_styled_heading(doc, '8.2 Serviceability', level=2)

    svc_items = [
        ('Deployment process', 'Automated build/deploy scripts (build.py, deploy.py). GitHub Actions triggers Docker build, pushes to ACR, deploys to Container Apps. UI-only overlay deploys in 30-60 seconds.'),
        ('Monitoring', 'Azure Monitor + Application Insights for infrastructure. Langfuse for AI-specific observability (prompt monitoring, token usage, cost tracking). Custom observability pipeline.'),
        ('Logging', 'Structured logging with Azure Log Analytics workspace. Per-tenant correlation via OpenTelemetry span attributes.'),
        ('Diagnostics', 'Health check endpoints on all containers. Pre-flight check scripts for staging and production environments. Test host containers for automated verification.'),
        ('Database management', 'Schema changes managed through application code (Cosmos DB is schema-flexible). Migration scripts in scripts/ directory with safety gates.'),
    ]
    for title, desc in svc_items:
        add_detail_item(doc, title, desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 9. MAINTAINABILITY
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '9. Maintainability', level=1)

    add_styled_heading(doc, '9.1 Strengths', level=2)
    strengths = [
        'Well-organized project structure with clear separation: src/ (backend), admin/ (3 SPAs), widget/, tests/, infrastructure/',
        'TenantScopedRepository base class eliminates boilerplate and enforces tenant isolation patterns',
        'Specification-first development workflow with knowledge database (2,092 specs, 11,055 tests tracked)',
        'Comprehensive CI/CD prevents regressions: 8 workflows including lint, test, build, deploy',
        'Shared component library (admin/shared/) reduces duplication across 3 admin SPAs',
        'Configuration-driven architecture (YAML schema registry) allows tenant customization without code changes',
        'Conventional commit format enables automated changelog generation',
    ]
    for item in strengths:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, '9.2 Concerns', level=2)
    concerns = [
        'Single-developer project with AI co-author — bus factor of 1. Knowledge transfer documentation is extensive but untested with new developers.',
        'Large codebase (2,213 files, 18 MB source) for a solo developer may become challenging to maintain as complexity grows.',
        'Three separate admin SPAs (standalone, Shopify, provider) with different Mantine versions (7.x and 8.x) may diverge over time.',
        'Dual AI governance system (CLAUDE.md + knowledge DB) adds operational overhead that may not scale with team growth.',
        'Some Azure service dependencies (Cosmos DB, Key Vault, NATS) are deeply coupled — migration to alternative providers would be costly.',
    ]
    for item in concerns:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 10. PERFORMANCE
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '10. Performance Considerations', level=1)

    add_styled_heading(doc, '10.1 Multi-Tenant Performance', level=2)

    perf_items = [
        ('AI Pipeline Latency', 'The multi-agent pipeline (intent classification, knowledge retrieval, response generation, critic review) introduces sequential latency. SSE streaming mitigates perceived latency by delivering tokens incrementally. Measured at approximately 5 seconds end-to-end, 69-80 tokens per response.'),
        ('Cosmos DB Performance', 'Partition key-based queries are O(1) at the database level. Cross-partition queries are avoided by design. Serverless mode handles bursty workloads; provisioned throughput prevents throttling at scale.'),
        ('Rate Limiting', 'Per-tenant rate limiting at 300 RPM (derived from ramp-to-overload testing: 1,380 RPM safe capacity per replica) provides noisy-neighbor protection.'),
        ('Container Scaling', 'HTTP-triggered and NATS-triggered auto-scaling independently scales each agent based on actual demand. Response Generator (most compute-intensive) scales up to 10 replicas.'),
        ('Caching', 'Redis caching for rate limiting and tenant configuration. Langfuse caching for AI observability. Additional caching opportunities exist for knowledge base embeddings.'),
    ]
    for title, desc in perf_items:
        add_detail_item(doc, title, desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 11. AZURE DEPLOYMENT
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '11. Azure Deployment Configuration', level=1)

    doc.add_paragraph(
        'Agent Red is deployed on Azure using Container Apps for compute, Cosmos DB for data, '
        'and a comprehensive set of managed services. Infrastructure is managed via Terraform (7 .tf files) '
        'with container deployments managed via Azure CLI / GitHub Actions.'
    )

    doc.add_picture(charts['azure_deploy'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 6, 'Current Agent Red Azure deployment showing 9 container apps, '
                       'managed services, and infrastructure-as-code components.')
    doc.add_paragraph()

    add_styled_heading(doc, '11.1 Container Apps Configuration', level=2)

    containers = [
        ('API Gateway', '0.5 CPU, 1 GB', '2-8 replicas', 'HTTP scaling', 'Main entry point. FastAPI. SSL termination.'),
        ('Intent Classifier', '0.5 CPU, 1 GB', '2-6 replicas', 'NATS scaling', 'GPT-4o-mini. Classifies customer intent.'),
        ('Knowledge Retrieval', '0.5 CPU, 1 GB', '2-6 replicas', 'NATS scaling', 'RAG pipeline with embeddings.'),
        ('Response Generator', '1.0 CPU, 2 GB', '2-10 replicas', 'NATS scaling', 'GPT-4o. Generates customer responses.'),
        ('Critic Supervisor', '0.5 CPU, 1 GB', '2-4 replicas', 'NATS scaling', 'Quality gate. Can retract responses.'),
        ('Escalation', '0.25 CPU, 0.5 GB', '1-3 replicas', 'NATS scaling', 'Human handoff orchestration.'),
        ('Analytics', '0.25 CPU, 0.5 GB', '1-2 replicas', 'NATS scaling', 'Metrics aggregation.'),
        ('Slim Gateway', '0.5 CPU, 1 GB', '2 (fixed)', 'Fixed', 'UI-only fast-rebuild overlay.'),
        ('NATS', '0.5 CPU, 1 GB', '2 (fixed)', 'Fixed', 'JetStream message bus.'),
    ]

    table = doc.add_table(rows=len(containers)+1, cols=5)
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Container', 'Resources', 'Replicas', 'Scaling', 'Purpose']):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(containers):
        for j, val in enumerate(row):
            table.cell(i+1, j).text = val
            for run in table.cell(i+1, j).paragraphs[0].runs:
                run.font.size = Pt(8)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 12. AZURE COST PROJECTIONS
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '12. Azure Cost Projections', level=1)

    doc.add_paragraph(
        'Agent Red\'s Azure costs are driven primarily by Azure OpenAI consumption (token-based), '
        'container compute, and Cosmos DB throughput. The following estimates model cost at three '
        'scaling tiers.'
    )

    doc.add_picture(charts['azure_costs'], width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 7, 'Estimated monthly Azure costs by scale tier. Azure OpenAI becomes the '
                       'dominant cost at scale, followed by container compute and Cosmos DB.')
    doc.add_paragraph()

    table = doc.add_table(rows=10, cols=4)
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(['Component', '10 Tenants', '100 Tenants', '1,000 Tenants']):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    cost_rows = [
        ('Container Apps (9 svcs)', '$220', '$900', '$5,500'),
        ('Cosmos DB', '$50', '$400', '$2,800'),
        ('Azure OpenAI', '$100', '$800', '$6,000'),
        ('Key Vault + Encryption', '$10', '$30', '$100'),
        ('NATS + Redis', '$30', '$80', '$300'),
        ('Monitoring (Langfuse+AI)', '$20', '$80', '$350'),
        ('Blob Storage + CDN', '$10', '$40', '$200'),
        ('Comms (Email/SMS)', '$5', '$50', '$400'),
        ('TOTAL', '$445/mo', '$2,380/mo', '$15,650/mo'),
    ]
    for i, row in enumerate(cost_rows):
        for j, val in enumerate(row):
            table.cell(i+1, j).text = val
            for run in table.cell(i+1, j).paragraphs[0].runs:
                run.font.size = Pt(9)
        if i == len(cost_rows) - 1:
            for j in range(4):
                for run in table.cell(i+1, j).paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Azure OpenAI costs are highly variable based on conversation volume and response length. '
        'Estimates assume average 3 conversations/tenant/day with ~500 tokens per conversation (input + output). '
        'Per-tenant cost at scale: approximately $15.65/tenant/month at 1,000 tenants. Reserved capacity and '
        'provisioned throughput discounts could reduce costs by 20-30% at scale.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 13. PRIORITIES
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, '13. Recommended Technical Priorities', level=1)

    doc.add_paragraph(
        'Agent Red\'s engineering maturity is high, shifting priorities from foundational concerns '
        'to market readiness, compliance, and partnership enablement.'
    )

    doc.add_picture(charts['roadmap'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 8, 'Recommended technical priority roadmap for Agent Red over 12 months.')
    doc.add_paragraph()

    priorities = [
        ('P1 — High (Months 1-6)', [
            ('Production Monitoring Enhancement', 'Extend Azure Monitor alerting rules. Add SLO dashboards (availability, latency p99, error rate). Implement on-call rotation alerting.'),
            ('Load/Stress Testing at Scale', 'Validate 100-tenant and 1,000-tenant performance using Locust load testing. Identify bottlenecks in NATS, Cosmos DB throughput, and OpenAI rate limits.'),
            ('SOC 2 Type II Preparation', 'Document security controls. Implement evidence collection automation. Engage third-party auditor. Target 6-month audit window.'),
            ('Public API Documentation', 'Publish OpenAPI-based developer documentation. Create integration quickstart guides. Enable self-serve API key provisioning.'),
        ]),
        ('P2 — Medium (Months 2-7)', [
            ('Multi-Region Disaster Recovery', 'Deploy secondary region (Canada East or East US 2). Configure Cosmos DB multi-region writes. Implement Azure Front Door for geographic routing and failover.'),
            ('French/Spanish Localization QA', 'Widget supports 8 languages but admin SPAs are English-only. Extend localization to admin interfaces. Professional translation review for fr-CA and es-MX.'),
            ('Mobile-Responsive Admin SPA', 'Optimize admin dashboards for tablet and mobile viewports. Priority: standalone admin SPA used by non-Shopify merchants.'),
        ]),
        ('P3 — Future (Months 5-12)', [
            ('Partner Integration SDK', 'Publish a partner SDK enabling CRM systems (like OrbaTech) to integrate bidirectionally. Include webhooks, OAuth, and event schemas.'),
            ('Self-Serve Onboarding', 'Automate tenant provisioning with self-serve signup, Stripe checkout, and guided setup wizard. Reduce time-to-value for new merchants.'),
            ('Marketplace / App Store', 'Enable third-party plugin agents via the existing plugin architecture. Publish agent marketplace for domain-specific AI capabilities.'),
        ]),
        ('Partnership Opportunity', [
            ('OrbaTech CRM Integration', 'Implement bidirectional data flow: CRM contact data enriches AI context, AI conversation insights flow back to CRM records. See Annex B for detailed architecture.'),
        ]),
    ]

    for group_title, items in priorities:
        add_styled_heading(doc, group_title, level=2)
        for title, desc in items:
            add_detail_item(doc, title, desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # ANNEX A
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, 'Annex A: Agent Red vs OrbaTech Recommended Azure Configuration', level=1)

    doc.add_paragraph(
        'This annex compares Agent Red\'s current Azure deployment against the recommended Azure '
        'configuration for OrbaTech CRM. The comparison highlights different architectural approaches '
        'to similar cloud challenges and identifies opportunities for shared infrastructure.'
    )

    doc.add_picture(charts['annex_comparison'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 9, 'Side-by-side comparison of Agent Red current and OrbaTech recommended Azure deployments.')
    doc.add_paragraph()

    add_styled_heading(doc, 'A.1 Key Architectural Differences', level=2)

    diffs = [
        ('Compute Model', 'Agent Red uses Azure Container Apps with 9 independently scaled microservices, enabling fine-grained resource allocation per agent. OrbaTech\'s recommended App Service PaaS model is simpler and more cost-effective for a monolithic .NET application.'),
        ('Database Strategy', 'Agent Red\'s Cosmos DB provides schema-flexible document storage with native partition-key tenant isolation, optimized for real-time AI conversations. OrbaTech\'s recommended Azure SQL with Elastic Pool is appropriate for structured CRM data with complex queries and reporting.'),
        ('Identity Approach', 'Agent Red implements custom identity with multiple authentication paths (JWT, API keys, widget keys) to support diverse integration scenarios (Shopify, standalone, embedded widget). OrbaTech\'s recommended Entra ID B2C provides managed consumer identity suitable for direct user authentication.'),
        ('CI/CD Maturity', 'Agent Red has a mature 8-workflow pipeline with 5-shard parallel testing, vulnerability scanning, and automated deployment. OrbaTech needs to build CI/CD from the ground up, but could leverage Agent Red\'s pipeline patterns as a template.'),
        ('Security Depth', 'Agent Red implements per-tenant envelope encryption with HSM-backed KEK, managed identities, and pre-auth rate limiting. OrbaTech\'s recommended configuration includes Row-Level Security and Key Vault, but has less encryption depth.'),
    ]
    for title, desc in diffs:
        add_detail_item(doc, title, desc)

    add_styled_heading(doc, 'A.2 Shared Infrastructure Opportunities', level=2)

    shared = [
        'Shared Azure Key Vault patterns for secrets management (both platforms use Key Vault)',
        'Common Azure Monitor / Application Insights workspace for cross-platform observability',
        'Shared Azure Communication Services instance for email and SMS delivery',
        'Common GitHub Actions workflow templates adaptable to both Python and .NET stacks',
        'Shared Azure Front Door configuration for unified SSL, WAF, and geographic routing',
        'Shared Container Registry (ACR) if OrbaTech eventually containerizes',
    ]
    for item in shared:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # ANNEX B
    # ══════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, 'Annex B: Agent Red + OrbaTech CRM Integration', level=1)

    doc.add_paragraph(
        'This annex describes a bidirectional integration architecture between Agent Red AI '
        'Customer Service and OrbaTech CRM. The integration enables CRM-enriched AI conversations '
        'and AI-informed customer relationship management.'
    )

    doc.add_picture(charts['integration'], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_caption(doc, 10, 'Bidirectional integration architecture between Agent Red AI and OrbaTech CRM.')
    doc.add_paragraph()

    add_styled_heading(doc, 'B.1 Agent Red to CRM (AI Insights)', level=2)

    ar_to_crm = [
        ('Conversation Summaries', 'AI-generated summaries of customer interactions are written to CRM contact records as activities, creating a complete interaction timeline accessible to sales teams.'),
        ('Sentiment Analysis', 'Real-time customer sentiment scores from AI conversations are attached to opportunity records, providing sales teams with early warning signals on at-risk deals.'),
        ('Quality Metrics', 'Per-interaction quality scores (from the critic agent) help sales managers identify training opportunities and customer satisfaction trends.'),
        ('Lead Qualification', 'AI-assessed purchase intent signals from customer conversations automatically create or update opportunity records in the CRM pipeline.'),
    ]
    for title, desc in ar_to_crm:
        add_detail_item(doc, title, desc)

    add_styled_heading(doc, 'B.2 CRM to Agent Red (Customer Context)', level=2)

    crm_to_ar = [
        ('Contact Context', 'Customer profiles, purchase history, and communication preferences from OrbaTech enrich Agent Red\'s AI context, enabling personalized responses.'),
        ('Opportunity Status', 'Active deals and pipeline stages inform Agent Red\'s conversation routing, ensuring high-value prospects receive priority handling and appropriate agent assignment.'),
        ('Activity History', 'Recent emails, calls, and meetings from CRM provide conversation context, reducing customer frustration from repeating information.'),
        ('Custom Fields', 'Tenant-specific custom fields from OrbaTech enable Agent Red to adapt its responses to each business\'s unique CRM taxonomy and domain vocabulary.'),
    ]
    for title, desc in crm_to_ar:
        add_detail_item(doc, title, desc)

    add_styled_heading(doc, 'B.3 Integration Technical Requirements', level=2)

    tech_reqs = [
        'OAuth 2.0 service-to-service authentication with client credentials flow',
        'Agent Red\'s existing integration framework (manifest-based registry) can register OrbaTech as a first-class adapter',
        'Webhook-based event notifications for real-time data synchronization',
        'Tenant mapping table linking Agent Red tenant identifiers to OrbaTech TenantId',
        'Data transformation layer leveraging Agent Red\'s normalized model pattern (SPEC-1762)',
        'Rate limiting and circuit breaker patterns (already implemented in Agent Red\'s adaptive rate limiter)',
        'Audit logging of all cross-platform data exchanges using Agent Red\'s append-only audit log',
    ]
    for item in tech_reqs:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ── Closing ──
    doc.add_page_break()
    add_styled_heading(doc, 'Disclaimer', level=1)
    doc.add_paragraph(
        'This report was prepared by Remaker Digital based on the Agent Red source code repository '
        'as of April 7, 2026. The findings and recommendations reflect the state of the codebase '
        'at that point in time. Technology assessments, cost estimates, and competitive comparisons '
        'are based on publicly available data and internal development records. Azure pricing estimates '
        'are approximate and subject to change.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u00a9 2026 Remaker Digital, a DBA of VanDusen & Palmeter, LLC. All rights reserved.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Save ──
    print(f"Saving to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    print(f"Report saved: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == '__main__':
    build_report()
