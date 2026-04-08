"""
Generate OrbaTech Technical Evaluation Report v2.
Addresses Codex review findings + adds expanded competitive matrix.

© 2026 Remaker Digital, a DBA of VanDusen & Palmeter, LLC. All rights reserved.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import tempfile

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "OrbaTech-Technical-Evaluation-Report.docx"
)

CHART_DIR = os.path.join(tempfile.gettempdir(), 'orbatech_charts')

# Chart file mapping
CHARTS = {
    'fig1': os.path.join(CHART_DIR, 'fig1_radar.png'),
    'fig2': os.path.join(CHART_DIR, 'fig2_architecture.png'),
    'fig3': os.path.join(CHART_DIR, 'fig3_security.png'),
    'fig4': os.path.join(CHART_DIR, 'fig4_features.png'),
    'fig5': os.path.join(CHART_DIR, 'fig5_azure.png'),
    'fig6': os.path.join(CHART_DIR, 'fig6_costs.png'),
    'fig7': os.path.join(CHART_DIR, 'fig7_roadmap.png'),
    'fig8': os.path.join(CHART_DIR, 'fig8_comparison.png'),
    'fig9': os.path.join(CHART_DIR, 'fig9_integration.png'),
    'fig10': os.path.join(CHART_DIR, 'fig10_summary.png'),
}

doc = Document()

# ── Style setup ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_chart(doc, chart_key, caption, width=Inches(6.0)):
    """Insert a chart image with caption."""
    path = CHARTS.get(chart_key)
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(f"[Chart not found: {chart_key}]")

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def add_severity_table(doc, rows):
    """Add security findings table with severity column."""
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Severity"
    table.rows[0].cells[1].text = "Finding"
    for p in table.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for p in table.rows[0].cells[1].paragraphs:
        for r in p.runs:
            r.bold = True
    for ri, (sev, finding) in enumerate(rows):
        table.rows[ri + 1].cells[0].text = sev
        table.rows[ri + 1].cells[1].text = finding
    return table


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("OrbaTech CRM")
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Technical Evaluation Report")
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Prepared by Remaker Digital\nApril 2026\nRevision 2")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run("CONFIDENTIAL")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Executive Summary",
    "2. Project Overview",
    "3. Technology Assessment",
    "4. Engineering Maturity",
    "5. Security Assessment",
    "6. Multi-Tenant Architecture",
    "7. Competitive Landscape",
    "8. Detailed Feature Comparison Matrix",
    "9. Reliability & Serviceability",
    "10. Maintainability",
    "11. Performance Considerations",
    "12. Recommended Azure Configuration",
    "13. Azure Cost Projections",
    "14. Recommended Technical Priorities",
    "Annex A: OrbaTech vs Agent Red Azure Deployment",
    "Annex B: OrbaTech + Agent Red Integration",
    "Annex C: Sources and Methodology",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    "This report provides an independent technical evaluation of the OrbaTech CRM platform, "
    "a Blazor-based customer relationship management system built on Microsoft .NET 8. The "
    "evaluation was conducted by Remaker Digital as part of a partnership assessment."
)
doc.add_paragraph(
    "OrbaTech is an early-stage product (approximately 3.5 months in development) built by "
    "a small team of three contributors. The platform targets small-to-medium businesses in "
    "the Canadian and US markets with core CRM functionality including contact management, "
    "deal pipeline tracking, email integration, and scheduling."
)
doc.add_heading('Overall Assessment', level=2)
doc.add_paragraph(
    "OrbaTech demonstrates solid technology choices (modern .NET 8 + Blazor) and a functional "
    "core architecture. The codebase includes a real test project with integration test "
    "infrastructure (xUnit, bUnit, Testcontainers, Respawn, FsCheck), EF Core global query "
    "filters for tenant isolation, and basic path traversal protection in file uploads. "
    "However, significant gaps exist in CI/CD automation, security hygiene, billing/entitlement "
    "infrastructure, compliance readiness, and cloud deployment that must be addressed before "
    "commercial production use. The platform has strong potential but requires structured "
    "engineering investment."
)
add_chart(doc, 'fig1',
    "Figure 1: Engineering maturity across eight dimensions (scale of 1-10). "
    "Tech Stack scores highest; CI/CD and Deployment require immediate attention.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Project Overview', level=1)
add_table(doc, ["Attribute", "Value"], [
    ["Repository", "github.com/QGolle/OrbaTech"],
    ["Description", "Blazor-based CRM for sales team deal management"],
    ["Primary Language", "C# (.NET 8.0)"],
    ["Project Age", "~3.5 months (since December 2025)"],
    ["Total Commits", "183+"],
    ["Contributors", "3 (1 primary, 2 supporting)"],
    ["License", "None (all rights reserved)"],
    ["Target Markets", "Canada and United States (EN primary, FR secondary, ES planned)"],
    ["Website", "orbatechcrm.com (source: external research, not repo content)"],
])
doc.add_paragraph(
    "The project demonstrates an active development cadence with daily commits from the "
    "primary contributor. The codebase comprises approximately 1.5 million bytes of source "
    "code across C#, HTML (Razor), CSS, JavaScript, and T-SQL."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Technology Assessment', level=1)
doc.add_heading('Technology Choices: Strong Foundation', level=2)
doc.add_paragraph(
    "The selection of .NET 8 with Blazor represents a modern, well-supported technology "
    "stack suitable for enterprise CRM applications. Microsoft provides long-term support "
    "for .NET 8 through November 2026, with a clear upgrade path to future LTS releases."
)
add_chart(doc, 'fig2',
    "Figure 2: Current OrbaTech application architecture showing the N-tier layout with "
    "Blazor hybrid rendering, dual ORM strategy, and application-level tenant isolation.")

doc.add_heading('3.1 Technology Stack Details', level=2)
add_table(doc, ["Layer", "Technology", "Assessment"], [
    ["Frontend", "Blazor Server + WebAssembly (Hybrid Rendering)",
     "Modern choice enabling SSR + interactive components. DevExpress UI suite provides enterprise-grade components."],
    ["Backend", "ASP.NET Core 8.0",
     "Industry-standard, high-performance web framework with excellent tooling and community support."],
    ["ORM", "Entity Framework Core 8.0 + Dapper 2.1",
     "Dual ORM strategy is pragmatic: EF Core for CRUD operations, Dapper for stored procedure calls and performance-critical paths."],
    ["Database", "SQL Server",
     "Appropriate for structured CRM data. Stored procedures (40+) and triggers (5) indicate investment in data-layer logic."],
    ["Authentication", "ASP.NET Core Identity",
     "Built-in framework with 2FA support. Adequate for current needs; should migrate to Entra ID B2C for cloud deployment."],
    ["UI Components", "DevExpress Blazor 25.2",
     "Commercial component suite. Provides grid, dashboard, scheduler, and rich text editor. Reduces development time but adds licensing costs."],
    ["Email", "MailKit (Windows Service)",
     "Standalone email processing service. Functions as spool watcher. Should be migrated to Azure Communication Services for cloud."],
    ["Testing", "xUnit + bUnit + Testcontainers + FsCheck",
     "Real integration test infrastructure with SQL Server containers, property-based testing, and Blazor component tests."],
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. ENGINEERING MATURITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Engineering Maturity', level=1)
doc.add_paragraph(
    "Engineering maturity was assessed across eight dimensions using a 1-10 scale. This "
    "assessment reflects the current state of the codebase and development practices as "
    "observed in the repository."
)

# CORRECTED per Codex finding #1: Testing score adjusted, description corrected
add_table(doc, ["Dimension", "Score", "Key Observations"], [
    ["Tech Stack", "8/10",
     "Modern .NET 8 + Blazor. Well-chosen dependencies. Current patch versions."],
    ["Architecture", "6/10",
     "Clean N-tier with repository/service pattern. Good separation of concerns."],
    ["Code Quality", "4/10",
     "No linting tools, no static analysis. Inconsistent nullable annotations. Commented-out code present."],
    ["Testing", "4/10",
     "Solution-included test project with xUnit, bUnit, Testcontainers (SQL Server), Respawn, and FsCheck. "
     "Concrete test files exist for services, repositories, and Blazor components. However, CI does not "
     "currently execute tests and no evidence of routine automated test runs is visible in workflows."],
    ["CI/CD", "1/10",
     "Only a csproj sanitization workflow (.github/workflows/clean-orbatech-csproj.yml). No build, test, or deploy automation."],
    ["Security", "3/10",
     "Identity/2FA implemented. However, credentials committed to repo. API keys in plaintext."],
    ["Documentation", "3/10",
     "Developer-oriented README covering user secrets, migrations, and test data. No architecture docs, no API guide, no deployment instructions."],
    ["Deployment", "2/10",
     "No containerization. No IaC. No cloud deployment configuration."],
])

doc.add_heading('4.1 Coding Standards', level=2)
doc.add_paragraph(
    "The codebase follows a consistent repository-service pattern with dependency injection, "
    "which demonstrates sound architectural thinking. However, several coding standard gaps "
    "were identified:"
)
# CORRECTED per Codex finding #4: Rephrased code review finding
standards_items = [
    "No .editorconfig file for consistent code formatting across contributors",
    "No Roslyn analyzers or StyleCop for compile-time code quality enforcement",
    "Inconsistent use of C# nullable reference type annotations",
    "Commented-out code in production files (e.g., HTTPS redirection)",
    "Commit messages contain spelling inconsistencies",
    "No repository-enforced review or CI policy artifacts are visible in the checked-in code "
    "(note: branch protection rules are repository settings, not code artifacts, so their "
    "presence or absence cannot be determined from the codebase alone)",
]
for item in standards_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. SECURITY ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Security Assessment', level=1)
p = doc.add_paragraph()
run = p.add_run("CRITICAL: Credentials Committed to Repository")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph(
    "Connection strings containing server names, a hardcoded API access token, and default "
    "administrator credentials were found committed to the repository. These must be removed "
    "immediately and rotated. All secrets should be managed via Azure Key Vault or .NET User Secrets."
)
add_chart(doc, 'fig3',
    "Figure 3: Security posture assessment across ten domains. Red and orange areas require "
    "immediate remediation.")

doc.add_heading('5.1 Findings Detail', level=2)
# CORRECTED per Codex finding #3: Split attachment finding into portability vs traversal
add_severity_table(doc, [
    ("CRITICAL",
     "Hardcoded connection strings in appsettings.json and appsettings.Production.json "
     "expose database server names. An API access token is committed in "
     "EmailService/appsettings.json."),
    ("CRITICAL",
     "Default administrator account seeded with the password \"Password123!\" \u2014 while "
     "gated to development environment, the seed logic and credentials are visible in the "
     "repository."),
    ("HIGH",
     "Database backup files (.bak) committed to the repository may contain production data "
     "and should be removed from version control history."),
    ("HIGH",
     "File attachment paths are hardcoded to Windows-specific local directories "
     "(appsettings.json, EmailService/appsettings.json), which prevents cloud deployment "
     "and creates a portability issue. Note: basic path traversal checks are visible in "
     "AttachmentService (Path.GetFileName stripping, canonical path resolution, upload "
     "directory containment check), so the risk is primarily filesystem portability, not "
     "traversal exploitation."),
    ("MEDIUM",
     "No API rate limiting observed. API key validation uses stored procedure "
     "sValidateApiKey.sql which compares plaintext tenant keys. Hash-based logic exists "
     "but is commented out."),
    ("MEDIUM",
     "HTTPS redirection is commented out (Program.cs:197), and AllowedHosts is set to "
     "wildcard (*) in all environments."),
])

doc.add_heading('5.2 Multi-Tenant Security Considerations', level=2)
doc.add_paragraph(
    "In a multi-tenant CRM environment, data isolation between tenants is a fundamental "
    "security requirement. OrbaTech implements application-level filtering using a TenantId "
    "column on each entity, with EF Core global query filters defined in ApplicationDbContext "
    "for Account, Contact, Opportunity, Layout, and TenantFieldType."
)
doc.add_paragraph(
    "However, the primary isolation risk is that repositories frequently bypass these "
    "filters: a search for IgnoreQueryFilters in the repository layer returns 47 hits. "
    "While some bypasses may be legitimate (admin operations, cross-tenant reporting), "
    "the volume suggests that isolation still depends heavily on developer discipline. "
    "The recommended mitigation is to implement SQL Server Row-Level Security (RLS) "
    "policies as a defense-in-depth measure and to audit every IgnoreQueryFilters "
    "call for necessity."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. MULTI-TENANT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Multi-Tenant Architecture', level=1)
doc.add_paragraph(
    "OrbaTech implements a shared-database, shared-schema multi-tenancy model. All tenants "
    "share a single SQL Server database with tenant isolation enforced through a TenantId "
    "foreign key on every business entity."
)

doc.add_heading('6.1 Current Implementation', level=2)
impl_items = [
    "TenantId consistently present on all major entities (Account, Contact, Opportunity, etc.)",
    "EF Core global query filters defined in ApplicationDbContext for automatic TenantId filtering "
    "on Account, Contact, Opportunity, Layout, and TenantFieldType",
    "Tenant claims injected at authentication time via custom ClaimsPrincipalFactory",
    "License tracking (LicenseCount, LicenseUsage) per tenant",
    "Dedicated test data with tenant-identifying names for leak detection",
    "SecurityContext service centralizes tenant resolution",
    "TenantManagementService and TenantService provide CRUD operations for tenant administration",
    "TenantRepository with stored procedures (gTenant.sql, gTenantAccounts.sql, AddTenantKey.sql)",
    "Customizable tenant field types (TenantFieldType entity) for per-tenant CRM customization",
]
for item in impl_items:
    doc.add_paragraph(item, style='List Bullet')

# CORRECTED per Codex finding #2: EF Core filters ARE present
doc.add_heading('6.2 Gaps and Recommendations', level=2)
add_table(doc, ["Gap", "Recommendation"], [
    ["Broad IgnoreQueryFilters() Usage (47 occurrences)",
     "Audit all 47 IgnoreQueryFilters() calls in the repository layer. Each bypass should "
     "be documented with a justification. Consider a custom Roslyn analyzer to flag new "
     "bypasses at compile time."],
    ["No Database-Level RLS",
     "Implement SQL Server Row-Level Security policies as a defense-in-depth mechanism "
     "that prevents cross-tenant access even if application code is bypassed."],
    ["Integer-Based TenantId",
     "Consider migrating to GUID-based tenant identifiers to prevent enumeration attacks "
     "and simplify multi-region database merges."],
    ["No Tenant Provisioning Automation",
     "Current provisioning is manual via seed data and admin tools. Implement automated "
     "tenant onboarding with self-service signup, configuration validation, license "
     "enforcement, and audit logging."],
    ["No Billing or Entitlement Enforcement",
     "No subscription management, usage metering, or billing integration exists. Implement "
     "a billing provider (Stripe, etc.) with entitlement checks at API boundaries."],
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. COMPETITIVE LANDSCAPE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Competitive Landscape', level=1)
doc.add_paragraph(
    "OrbaTech enters a mature and competitive CRM market. The following comparison positions "
    "OrbaTech against established players serving the SMB segment in North America. This "
    "section uses two sources: repository analysis (features visible in checked-in code) and "
    "external market research (competitor pricing, certifications, integration counts). "
    "External claims are marked with \u2020 where applicable."
)
add_chart(doc, 'fig4',
    "Figure 4: Feature capability comparison across twelve CRM dimensions. OrbaTech shows "
    "strength in core CRM functions but has gaps in mobile, API, AI, and localization.",
    width=Inches(6.5))

doc.add_heading('7.1 Competitor Profiles', level=2)
add_table(doc, ["Competitor", "Pricing\u2020", "Market Position\u2020"], [
    ["Salesforce", "$25\u2013$300/user/mo",
     "Market leader. Most comprehensive feature set. Strong enterprise adoption. "
     "Often perceived as complex and expensive for SMBs."],
    ["HubSpot CRM", "Free\u2013$1,200/mo",
     "Leading SMB CRM with generous free tier. Strong marketing automation and "
     "content management. Growing enterprise presence."],
    ["Zoho CRM", "$14\u2013$52/user/mo",
     "Cost-effective with broad feature coverage. 28 languages. Strong in "
     "price-sensitive SMB segments."],
    ["Pipedrive", "$14\u2013$99/user/mo",
     "Sales-focused with intuitive pipeline management. Popular with small sales "
     "teams. Limited customization."],
    ["Freshsales (Freshworks)", "$9\u2013$59/user/mo",
     "AI-powered CRM with built-in phone and email. Growing mid-market presence. "
     "Part of broader Freshworks suite."],
    ["Monday Sales CRM", "$12\u2013$28/user/mo",
     "Visual, work-management-centric CRM. Strong customization via low-code "
     "automations. Appeals to non-traditional CRM users."],
])

doc.add_heading('7.2 OrbaTech Differentiation Opportunities', level=2)
diff_items = [
    "Canadian data residency (compliance with PIPEDA) \u2014 a differentiator for Canadian businesses",
    "Bilingual (EN/FR) native support \u2014 required for Quebec market, underserved by most US-based CRMs",
    "Simpler pricing and onboarding for micro-businesses (1\u201310 users)",
    "Integration with AI customer service (Agent Red partnership potential)",
    "Industry-specific customization for Canadian market verticals",
    ".NET/Blazor technology stack appeals to Microsoft-ecosystem enterprises",
]
for item in diff_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. DETAILED FEATURE COMPARISON MATRIX (NEW SECTION)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Detailed Feature Comparison Matrix', level=1)
doc.add_paragraph(
    "This section provides a detailed comparison across six critical enterprise SaaS "
    "dimensions requested for partnership evaluation. OrbaTech capabilities are assessed "
    "from repository evidence; competitor capabilities are assessed from public documentation "
    "and market research (marked \u2020). Ratings use a five-level scale: "
    "\u2588\u2588\u2588\u2588\u2588 Mature, \u2588\u2588\u2588\u2588\u25a1 Good, "
    "\u2588\u2588\u2588\u25a1\u25a1 Developing, \u2588\u2588\u25a1\u25a1\u25a1 Basic, "
    "\u2588\u25a1\u25a1\u25a1\u25a1 Minimal/None."
)

# 8.1 Tenant Provisioning
doc.add_heading('8.1 Tenant Provisioning', level=2)
doc.add_paragraph(
    "Tenant provisioning covers how new customer organizations are created, configured, and "
    "activated on the platform. Mature provisioning includes self-service signup, automated "
    "configuration, trial management, and programmatic API access."
)
add_table(doc, ["Platform", "Self-Service Signup", "Admin Provisioning", "API Provisioning",
                "Trial Management", "Assessment"], [
    ["OrbaTech",
     "None. No public registration endpoint.",
     "Manual via seed data (SeedInitialAdminAccount.cs) and admin test data dashboard (/test-data-dashboard). "
     "TenantManagementService provides CRUD operations.",
     "None. No provisioning API exposed.",
     "None. No trial/freemium tier.",
     "Basic. Manual provisioning only."],
    ["Salesforce\u2020",
     "Yes (developer edition). Enterprise via sales.",
     "Full admin console. Sandbox provisioning.",
     "Yes. Salesforce Metadata API, Tooling API.",
     "30-day trials. Sandbox environments.",
     "Mature. Full lifecycle automation."],
    ["HubSpot\u2020",
     "Yes. Free-tier self-service.",
     "Full admin portal. Team/role setup.",
     "Yes. HubSpot API for account management.",
     "14-day trials for paid tiers.",
     "Mature. Freemium-driven onboarding."],
    ["Zoho\u2020",
     "Yes. Free-tier self-service.",
     "Admin console with org hierarchy.",
     "Yes. Zoho REST APIs.",
     "15-day trials.",
     "Good. Self-service with admin tools."],
    ["Pipedrive\u2020",
     "Yes. Self-service signup.",
     "Admin settings for team management.",
     "Yes. REST API.",
     "14-day trial.",
     "Good. Sales-team focused."],
    ["Freshsales\u2020",
     "Yes. Free-tier self-service.",
     "Admin console.",
     "Yes. REST API.",
     "21-day trial.",
     "Good. Simple onboarding flow."],
    ["Monday CRM\u2020",
     "Yes. Self-service.",
     "Admin dashboard.",
     "Yes. GraphQL API.",
     "14-day trial.",
     "Good. Visual onboarding."],
])

# 8.2 Metering, Billing, and Entitlements
doc.add_heading('8.2 Metering, Billing, and Entitlements', level=2)
doc.add_paragraph(
    "This dimension covers usage tracking, subscription management, billing integration, "
    "and feature entitlement enforcement. It is critical for commercial SaaS viability."
)
add_table(doc, ["Platform", "Billing Integration", "Usage Metering", "Entitlement Enforcement",
                "Subscription Mgmt", "Assessment"], [
    ["OrbaTech",
     "None. No payment processor integration. BillingCodeLength field is a CRM data field, not billing infrastructure.",
     "None. No usage tracking or metering.",
     "LicenseCount and LicenseUsage fields exist per tenant but no enforcement logic observed at API boundaries.",
     "None. No subscription lifecycle management.",
     "Minimal. License fields exist but are not enforced."],
    ["Salesforce\u2020",
     "Salesforce Billing product (CPQ). Stripe integration available.",
     "API call limits, storage quotas, feature usage dashboards.",
     "Edition-based (Essentials/Professional/Enterprise/Unlimited) with per-feature flags.",
     "Full subscription lifecycle via Salesforce Billing or AppExchange partners.",
     "Mature. Comprehensive billing and entitlement platform."],
    ["HubSpot\u2020",
     "Stripe integration for Payments tool. HubSpot Commerce Hub.",
     "Marketing contacts, API calls, storage tracked per plan.",
     "Hub-based (Marketing/Sales/Service) and tier-based (Starter/Pro/Enterprise) gating.",
     "Self-service upgrade/downgrade. Annual and monthly plans.",
     "Mature. Integrated commerce capabilities."],
    ["Zoho\u2020",
     "Zoho Billing (Zoho Subscriptions). Stripe, PayPal integrations.",
     "API limits, storage, record count tracked.",
     "Edition-based with feature flags. Zoho One bundling.",
     "Zoho Subscriptions product for full lifecycle.",
     "Mature. Native billing product in ecosystem."],
    ["Pipedrive\u2020",
     "Stripe via marketplace. Basic invoicing.",
     "API rate limits. Deal/contact limits per plan.",
     "Plan-based feature gating (Essential/Advanced/Professional/Power/Enterprise).",
     "Self-service portal.",
     "Good. Standard SaaS billing model."],
    ["Freshsales\u2020",
     "Freshworks marketplace integrations.",
     "API limits, contact/account limits per plan.",
     "Plan-based (Free/Growth/Pro/Enterprise) with add-on modules.",
     "Self-service with Freshworks billing.",
     "Good. Modular entitlement model."],
    ["Monday CRM\u2020",
     "Marketplace integrations. No native billing.",
     "Seat-based pricing. Automation/integration limits per plan.",
     "Plan-based (Basic/Standard/Pro/Enterprise) with feature gating.",
     "Self-service upgrade/downgrade.",
     "Good. Seat-based with plan gating."],
])

# 8.3 Compliance
doc.add_heading('8.3 Compliance (SOC-2, HIPAA, GDPR, etc.)', level=2)
doc.add_paragraph(
    "Compliance readiness covers security certifications, regulatory frameworks, data "
    "protection mechanisms, audit capabilities, and privacy tooling."
)
add_table(doc, ["Platform", "SOC-2", "HIPAA", "GDPR", "Other Certifications",
                "Audit Logging", "Data Export/Deletion", "Assessment"], [
    ["OrbaTech",
     "No.", "No.", "No.",
     "None. No compliance certifications or privacy tooling visible in repo.",
     "No audit logging infrastructure. Basic API security (APISecurityContext.cs).",
     "No data export or deletion tooling for GDPR/PIPEDA compliance.",
     "None. No compliance readiness."],
    ["Salesforce\u2020",
     "Yes (Type II).", "Yes (BAA available).", "Yes.",
     "ISO 27001, ISO 27018, FedRAMP, PCI DSS, C5.",
     "Field-level audit trail. Setup audit trail. Event monitoring.",
     "Data export, GDPR deletion tools, privacy center.",
     "Mature. Industry-leading certifications."],
    ["HubSpot\u2020",
     "Yes (Type II).", "No (limited data handling).", "Yes.",
     "ISO 27001. Privacy Shield successor (DPF).",
     "Activity logging per contact/deal. Login audit logs.",
     "GDPR tools (consent, deletion, export). Cookie consent.",
     "Good. Strong GDPR tooling."],
    ["Zoho\u2020",
     "Yes (Type II).", "Yes (BAA for Zoho One).", "Yes.",
     "ISO 27001, ISO 27017, ISO 27018.",
     "Admin audit logs. Field-level audit for CRM Plus.",
     "GDPR module (consent, erasure, portability).",
     "Mature. Broad certification coverage."],
    ["Pipedrive\u2020",
     "Yes (Type II).", "No.", "Yes.",
     "ISO 27001.",
     "Changelog/activity logs.",
     "GDPR tools (consent, export, deletion).",
     "Good. Standard compliance for SMB CRM."],
    ["Freshsales\u2020",
     "Yes (Type II).", "Yes (with Freshworks BAA).", "Yes.",
     "ISO 27001.",
     "Audit logs for admin actions.",
     "GDPR compliance tools.",
     "Good. Growing certification portfolio."],
    ["Monday CRM\u2020",
     "Yes (Type II).", "Yes (Enterprise plan).", "Yes.",
     "ISO 27001, ISO 27018.",
     "Audit log (Enterprise plan).",
     "GDPR tools, data export.",
     "Good. Enterprise-tier compliance focus."],
])

# 8.4 Integrations
doc.add_heading('8.4 Integrations with Complementary Services', level=2)
doc.add_paragraph(
    "Integration breadth and depth determine how well a CRM fits into existing business "
    "workflows. This covers native integrations, marketplace/ecosystem size, and key "
    "partnership categories."
)
add_table(doc, ["Platform", "Native Integrations", "Marketplace", "Email", "Calendar",
                "Phone/VoIP", "Social", "Payment", "Assessment"], [
    ["OrbaTech",
     "2: Salesforce import (stored procs for Account/Contact/Opportunity), MailKit email service.",
     "None.",
     "MailKit outbound only (Windows service spool watcher).",
     "Scheduler component visible (DevExpress) but no external calendar sync.",
     "None.",
     "None.",
     "None.",
     "Minimal. Email sending and SF import only."],
    ["Salesforce\u2020",
     "100+ native.",
     "AppExchange: 7,000+ apps.",
     "Gmail, Outlook, native email.",
     "Google Calendar, Outlook.",
     "Salesforce Voice, CTI adapters.",
     "Social Studio, LinkedIn Sales Nav.",
     "Stripe, PayPal via AppExchange.",
     "Mature. Largest ecosystem."],
    ["HubSpot\u2020",
     "50+ native.",
     "App Marketplace: 1,500+ apps.",
     "Gmail, Outlook, native email.",
     "Google Calendar, Outlook, native meetings.",
     "HubSpot Calling, Aircall, RingCentral.",
     "LinkedIn, Facebook, Instagram, Twitter.",
     "Stripe, native Payments tool.",
     "Mature. Strong partner ecosystem."],
    ["Zoho\u2020",
     "40+ native (Zoho ecosystem).",
     "Zoho Marketplace: 1,000+ extensions.",
     "Zoho Mail, Gmail, Outlook.",
     "Zoho Calendar, Google, Outlook.",
     "Zoho Telephony, Twilio, RingCentral.",
     "Zoho Social, LinkedIn.",
     "Zoho Books, Stripe, PayPal.",
     "Good. Deep ecosystem integration."],
    ["Pipedrive\u2020",
     "30+ native.",
     "Marketplace: 500+ apps.",
     "Gmail, Outlook, native email.",
     "Google Calendar, Outlook.",
     "Pipedrive Calling, Aircall.",
     "Limited social integrations.",
     "Stripe via marketplace.",
     "Good. Focused on sales tools."],
    ["Freshsales\u2020",
     "30+ native.",
     "Freshworks Marketplace: 1,000+ apps.",
     "Gmail, Outlook, native email.",
     "Google Calendar, Outlook.",
     "Built-in phone (Freshcaller).",
     "Facebook, Instagram.",
     "Marketplace partners.",
     "Good. Built-in phone is differentiator."],
    ["Monday CRM\u2020",
     "50+ native.",
     "Apps Marketplace: 200+ apps.",
     "Gmail, Outlook.",
     "Google Calendar, Outlook.",
     "Aircall, Twilio via marketplace.",
     "Limited.",
     "Marketplace partners.",
     "Developing. Growing ecosystem."],
])

# 8.5 Documentation
doc.add_heading('8.5 Documentation', level=2)
doc.add_paragraph(
    "Documentation quality affects developer onboarding, customer adoption, partner "
    "integration, and operational sustainability."
)
add_table(doc, ["Platform", "API Documentation", "Developer Guides", "Architecture Docs",
                "User Guides", "Knowledge Base", "Assessment"], [
    ["OrbaTech",
     "None. No API documentation. Stored procedures have inline comments only.",
     "README covers: GitHub Actions csproj cleanup, .NET user secrets setup, EF Core migrations, "
     "test data generation. Scattered readme.txt files in DB/ER/, DB/SP/, EmailService/.",
     "None. No architecture decision records, no design documents.",
     "None. No end-user documentation.",
     "None.",
     "Basic. Developer README only."],
    ["Salesforce\u2020",
     "Comprehensive. REST, SOAP, Metadata, Bulk, Streaming APIs. OpenAPI specs.",
     "Trailhead learning platform. Developer documentation portal.",
     "Architecture center with design patterns, data modeling guides.",
     "Extensive help documentation. Trailhead modules.",
     "Salesforce Help. Community forums. Stack Exchange.",
     "Mature. Industry-leading documentation."],
    ["HubSpot\u2020",
     "Full REST API docs with interactive explorer. OpenAPI specs.",
     "Developer portal with guides, tutorials, and sample code.",
     "Architecture overview for integrations. CMS architecture docs.",
     "Knowledge Base with role-based guides.",
     "HubSpot Academy. Community forums.",
     "Mature. Developer-friendly with education platform."],
    ["Zoho\u2020",
     "REST API documentation for all products.",
     "Developer console. SDK documentation for multiple languages.",
     "Solution architecture guides.",
     "Product documentation per module.",
     "Zoho Desk knowledge base. Community forums.",
     "Good. Comprehensive but fragmented across products."],
    ["Pipedrive\u2020",
     "REST API v1/v2 with OpenAPI. Webhooks documented.",
     "Developer documentation with code samples.",
     "Limited architecture docs.",
     "Help center with guides.",
     "Community forum. Academy courses.",
     "Good. API-first documentation approach."],
    ["Freshsales\u2020",
     "REST API documentation.",
     "Developer portal with integration guides.",
     "Limited public architecture docs.",
     "Product help documentation.",
     "Freshworks Academy.",
     "Good. Standard SaaS documentation."],
    ["Monday CRM\u2020",
     "GraphQL API documentation.",
     "Developer portal with app framework docs.",
     "Limited.",
     "Help center with guides and videos.",
     "Monday Community.",
     "Developing. GraphQL API is well-documented."],
])

# 8.6 Design and Architecture Documents
doc.add_heading('8.6 Design and Architecture Documents', level=2)
doc.add_paragraph(
    "Architecture documentation reflects engineering maturity and facilitates onboarding, "
    "security review, and partnership evaluation."
)
add_table(doc, ["Platform", "ADRs / Design Docs", "Data Model Docs", "ER Diagrams",
                "Deployment Docs", "Security Architecture", "Assessment"], [
    ["OrbaTech",
     "None. No architecture decision records. No design documents in repo.",
     "No explicit data model documentation. Schema is inferred from EF Core models and stored procedures.",
     "One Erwin model file (DB/ER/OrbaTechCloudChamp.erwin). No exported diagrams.",
     "None. No deployment instructions or runbooks.",
     "None. No security architecture documentation.",
     "Minimal. ER model file exists but is tooling-specific."],
    ["Salesforce\u2020",
     "Architecture Center with design patterns. Well Architected guidance.",
     "ERDs for standard and custom objects. Schema Builder tool.",
     "Schema Builder visual tool.",
     "Deployment guides for orgs, sandboxes, and managed packages.",
     "Security architecture documentation. Trust site (trust.salesforce.com).",
     "Mature. Comprehensive architecture guidance."],
    ["HubSpot\u2020",
     "Public engineering blog with architectural decisions.",
     "Object schema documentation. Custom objects API.",
     "Not publicly available.",
     "CMS deployment guides. CI/CD integration docs.",
     "Security practices documentation.",
     "Good. Engineering blog supplements formal docs."],
    ["Zoho\u2020",
     "Limited public ADRs.",
     "Module relationship documentation.",
     "Not publicly available.",
     "On-premises deployment guides (Zoho One).",
     "Security and compliance center.",
     "Developing. Operational rather than architectural."],
    ["Pipedrive\u2020",
     "Engineering blog with some architectural content.",
     "API reference describes data model.",
     "Not publicly available.",
     "SaaS-only. No deployment docs needed.",
     "Security practices page.",
     "Developing. API-driven data model docs."],
    ["Freshsales\u2020",
     "Limited.",
     "API reference.",
     "Not publicly available.",
     "SaaS-only.",
     "Security practices documentation.",
     "Developing."],
    ["Monday CRM\u2020",
     "Limited.",
     "GraphQL schema documentation.",
     "Not publicly available.",
     "SaaS-only.",
     "Security page.",
     "Developing. GraphQL schema is self-documenting."],
])

doc.add_heading('8.7 Summary Assessment Matrix', level=2)
add_chart(doc, 'fig10',
    "Figure 10: Enterprise readiness heatmap and total scores across seven CRM platforms.",
    width=Inches(6.5))
doc.add_paragraph(
    "Consolidated readiness scores across all six dimensions. Scale: 5 = Mature, "
    "4 = Good, 3 = Developing, 2 = Basic, 1 = Minimal/None."
)
add_table(doc, ["Platform", "Provisioning", "Billing", "Compliance", "Integrations",
                "Documentation", "Architecture Docs", "TOTAL (out of 30)"], [
    ["OrbaTech", "2", "1", "1", "1", "2", "1", "8"],
    ["Salesforce\u2020", "5", "5", "5", "5", "5", "5", "30"],
    ["HubSpot\u2020", "5", "5", "4", "5", "5", "4", "28"],
    ["Zoho\u2020", "4", "5", "5", "4", "4", "3", "25"],
    ["Pipedrive\u2020", "4", "4", "4", "4", "4", "3", "23"],
    ["Freshsales\u2020", "4", "4", "4", "4", "4", "3", "23"],
    ["Monday CRM\u2020", "4", "4", "4", "3", "3", "3", "21"],
])
doc.add_paragraph(
    "\u2020 Competitor assessments based on publicly available documentation and market research "
    "as of April 2026, not repository analysis."
)
doc.add_paragraph(
    "OrbaTech's total score of 8/30 reflects its early-stage status. The gap is expected "
    "for a 3.5-month-old product built by a three-person team. The key question for "
    "partnership evaluation is the team's capacity and timeline to close these gaps."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. RELIABILITY & SERVICEABILITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Reliability & Serviceability', level=1)
doc.add_heading('9.1 Reliability', level=2)
doc.add_paragraph(
    "Reliability measures the system's ability to operate without failure. The current state "
    "presents several reliability concerns:"
)
reliability_items = [
    "No automated testing in CI: Test infrastructure exists (xUnit, bUnit, Testcontainers) but "
    "tests are not executed in CI workflows. Regressions may go undetected.",
    "No health check endpoints: No visible health or readiness probe endpoints for load balancer "
    "or orchestrator integration.",
    "Single-server deployment model: No horizontal scaling, load balancing, or failover configuration observed.",
    "Windows service dependency: The EmailService runs as a Windows service tied to a specific "
    "machine, creating a single point of failure.",
    "No error tracking: No integration with error tracking services (e.g., Application Insights, "
    "Sentry) was observed.",
]
for item in reliability_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9.2 Serviceability', level=2)
doc.add_paragraph(
    "Serviceability measures how easily the system can be maintained, updated, and diagnosed "
    "in production."
)
service_items = [
    "Deployment process: No documented deployment process. No deployment automation.",
    "Logging: Standard ASP.NET Core logging configured. No structured logging or centralized "
    "log aggregation observed.",
    "Diagnostics: No diagnostic endpoints, no performance counters, no distributed tracing.",
    "Database management: EF Core migrations present but no automated migration strategy for "
    "zero-downtime deployments.",
    "Configuration management: Environment-specific appsettings files exist but contain "
    "hardcoded values that should be externalized.",
]
for item in service_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. MAINTAINABILITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Maintainability', level=1)
doc.add_heading('10.1 Strengths', level=2)
strengths = [
    "Consistent repository-service pattern with interface-based dependency injection",
    "Clear project structure with logical separation of concerns (Models, Repositories, "
    "Services, Components)",
    "Shared class library for DTOs and interfaces promotes code reuse between server and "
    "client projects",
    "Entity Framework Core migrations provide version-controlled schema management",
    "Scoped CSS styles per component prevent style conflicts",
    "Real test infrastructure with integration tests (Testcontainers for SQL Server), "
    "property-based tests (FsCheck), and Blazor component tests (bUnit)",
]
for item in strengths:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.2 Concerns', level=2)
concerns = [
    "Dual ORM strategy (EF Core + Dapper + stored procedures) creates three places where "
    "data access logic can live, increasing cognitive load",
    "No automated code quality enforcement \u2014 relies entirely on developer discipline",
    "Database backup files and migration archives committed to the repository inflate "
    "repository size",
    "No contribution guidelines or coding standards document for onboarding new developers",
    "DevExpress commercial licensing creates vendor lock-in for UI components",
]
for item in concerns:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. PERFORMANCE CONSIDERATIONS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Performance Considerations', level=1)
doc.add_heading('11.1 Multi-Tenant Performance', level=2)
doc.add_paragraph(
    "In a multi-tenant environment, performance isolation between tenants is essential "
    "to prevent one tenant's heavy usage from degrading service for others."
)
perf_items = [
    "Query Performance: The dual ORM approach provides good flexibility for performance "
    "optimization. Stored procedures can be tuned independently of application code.",
    "Connection Pooling: Standard ADO.NET connection pooling. At 100+ tenants, pool "
    "exhaustion may occur without partitioning or per-tenant limits.",
    "Caching: No caching layer observed. Adding Redis or in-memory caching for tenant "
    "configurations and reference data would improve response times.",
    "Blazor Server SignalR: Persistent SignalR connection per user. At 1,000+ concurrent "
    "users, significant memory and connection pressure.",
    "Blazor WebAssembly: Hybrid rendering offloads interactive components to client. "
    "Positive architectural choice for reducing server load.",
]
for item in perf_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. RECOMMENDED AZURE CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('12. Recommended Azure Configuration', level=1)
doc.add_paragraph(
    "The following Azure architecture is recommended for OrbaTech CRM targeting production "
    "deployment in the Canadian and US markets. The design emphasizes data residency "
    "compliance (Canada Central primary region), security best practices, and cost-effective "
    "scaling."
)
add_chart(doc, 'fig5',
    "Figure 5: Recommended Azure architecture for OrbaTech CRM production deployment.",
    width=Inches(6.5))

doc.add_heading('12.1 Core Services', level=2)
services = [
    ("Azure App Service (P1v3):",
     "Primary compute for Blazor. 2 vCPUs, 8 GB RAM, deployment slots for zero-downtime "
     "updates. Auto-scale 2\u20136 instances."),
    ("Azure SQL Database:",
     "Start with Standard S2 (50 DTU). Migrate to Elastic Pool at 100+ tenants."),
    ("Azure Front Door:",
     "Global load balancing with SSL, WAF, and geographic routing (Canada Central + East US)."),
    ("Azure Blob Storage:",
     "Replace local filesystem for attachments. Standard LRS with lifecycle management."),
    ("Azure Key Vault:",
     "Centralized secrets management with managed identity authentication."),
    ("Azure Entra ID B2C:",
     "Cloud-native identity replacing ASP.NET Core Identity. SSO, MFA, social login."),
    ("Azure Communication Services:",
     "Replace Windows-based EmailService. Email, SMS, push notifications."),
    ("Azure Monitor + Application Insights:",
     "Full-stack monitoring with distributed tracing, live metrics, AI anomaly detection."),
    ("Azure Redis Cache:",
     "Caching for tenant config, session state, reference data. Basic C0 at startup."),
]
for title, desc in services:
    p = doc.add_paragraph()
    run = p.add_run(title + " ")
    run.bold = True
    p.add_run(desc)

doc.add_heading('12.2 Security Configuration', level=2)
sec_items = [
    "Managed Identities for all service-to-service authentication",
    "Private Endpoints for SQL Database, Key Vault, and Redis",
    "SQL Server Row-Level Security (RLS) policies for tenant isolation",
    "Azure RBAC with least-privilege access",
    "Azure Policy for compliance enforcement (data residency, encryption)",
]
for item in sec_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 13. AZURE COST PROJECTIONS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('13. Azure Cost Projections', level=1)
doc.add_paragraph(
    "The following cost estimates are based on Azure Canada Central region pricing. "
    "Costs are projected across three scaling tiers representing typical growth stages "
    "for an SMB CRM platform."
)
add_chart(doc, 'fig6',
    "Figure 6: Estimated monthly Azure costs across three scaling tiers.",
    width=Inches(6.5))
# CORRECTED per Codex finding #5: Explicit source attribution
doc.add_paragraph(
    "Source: Azure Pricing Calculator (pricing.azure.com), April 2026 rates. These are "
    "external research estimates, not derived from the OrbaTech repository.",
    style='Intense Quote'
)
add_table(doc, ["Component", "10 Tenants", "100 Tenants", "1,000 Tenants"], [
    ["Compute (App Service)", "$50", "$300", "$1,800"],
    ["SQL Database", "$30", "$250", "$1,500"],
    ["Storage & CDN", "$5", "$30", "$200"],
    ["Identity (Entra ID B2C)", "$0", "$50", "$500"],
    ["Monitoring & Logging", "$10", "$50", "$200"],
    ["Networking & DNS", "$10", "$30", "$100"],
    ["Email Service", "$0", "$20", "$100"],
    ["Backup & DR", "$5", "$50", "$300"],
    ["TOTAL", "$110/mo", "$780/mo", "$4,700/mo"],
])
doc.add_paragraph(
    "Note: Assumes Azure Reserved Instances (1-year, ~35% savings) and standard support. "
    "DevExpress licensing (~$2,400/year per developer seat) not included. Per-tenant cost "
    "at scale: approximately $4.70/tenant/month at 1,000 tenants."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 14. RECOMMENDED TECHNICAL PRIORITIES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('14. Recommended Technical Priorities', level=1)
add_chart(doc, 'fig7',
    "Figure 7: Recommended technical priority roadmap over 12 months.",
    width=Inches(6.5))

doc.add_heading('P0 \u2014 Critical (Months 1\u20133)', level=2)
p0_items = [
    "Security Remediation: Remove all committed credentials from repository history "
    "(git filter-repo). Rotate all exposed secrets. Implement Azure Key Vault. Enable HTTPS.",
    "CI/CD Pipeline: GitHub Actions for build, test, and deployment. Vulnerability scanning "
    "(Dependabot). Branch protection with PR reviews.",
    "Automated Testing in CI: Connect existing xUnit/bUnit/Testcontainers test suite to "
    "CI pipeline. Add multi-tenant data isolation tests. Target 60% coverage.",
    "Billing and Entitlements: Implement billing provider integration (Stripe recommended). "
    "Enforce LicenseCount limits at API boundaries. Add usage metering.",
]
for item in p0_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('P1 \u2014 High (Months 2\u20136)', level=2)
p1_items = [
    "Multi-Tenant RLS: SQL Server Row-Level Security. Audit all 47 IgnoreQueryFilters() "
    "calls. Cross-tenant penetration testing.",
    "Containerization: Dockerfile for main app and EmailService. docker-compose for local "
    "development. Azure Container Apps or App Service deployment.",
    "Azure Deployment: Infrastructure as Code (Bicep/Terraform). Staging + production "
    "environments. Deployment slots for zero-downtime releases.",
    "Compliance Foundations: GDPR/PIPEDA data export and deletion tooling. Audit logging "
    "infrastructure. SOC-2 readiness assessment.",
    "API Authentication: OAuth 2.0 / OpenID Connect. Rate limiting. API versioning.",
    "Self-Service Tenant Provisioning: Public registration endpoint. Trial management. "
    "Automated configuration validation.",
]
for item in p1_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('P2 \u2014 Medium (Months 5\u20138)', level=2)
p2_items = [
    "Localization (EN/FR): Resource-based localization. French Canadian (fr-CA) translations. "
    "Plan for Spanish (es).",
    "Performance Optimization: Redis caching. Connection pool management. Stored procedure "
    "optimization. Response compression.",
    "Documentation: API documentation (OpenAPI/Swagger). Architecture decision records. "
    "Deployment runbooks. User guides.",
    "Integration Expansion: Calendar sync (Google/Outlook). VoIP integration. Social media "
    "connectors.",
]
for item in p2_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('P3 \u2014 Future (Months 7\u201312)', level=2)
p3_items = [
    "Mobile Responsive: Optimize Blazor for mobile. Consider PWA.",
    "AI/Automation: Predictive lead scoring, email templates, workflow automation.",
    "Agent Red Integration: Bidirectional data flow for AI-powered customer service "
    "enrichment (see Annex B).",
    "SOC-2 Certification: Formal audit engagement after controls implemented.",
]
for item in p3_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ANNEX A
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Annex A: OrbaTech vs Agent Red Azure Deployment', level=1)
doc.add_paragraph(
    "This annex compares the recommended Azure configuration for OrbaTech CRM against "
    "Remaker Digital's Agent Red platform deployment."
)
add_chart(doc, 'fig8',
    "Figure 8: Side-by-side OrbaTech and Agent Red Azure deployments.",
    width=Inches(6.5))

doc.add_heading('A.1 Key Architectural Differences', level=2)
arch_diffs = [
    "Compute Model: OrbaTech recommended for Azure App Service (PaaS). Agent Red uses "
    "Azure Container Apps for microservice flexibility (8+ containers).",
    "Database Strategy: OrbaTech uses SQL Server for structured CRM data. Agent Red uses "
    "Cosmos DB for globally distributed document storage.",
    "Identity: OrbaTech should adopt Entra ID B2C. Agent Red implements custom identity "
    "with widget keys and API keys for embedded SaaS.",
    "Tenant Isolation: Both use application-level filtering. Agent Red adds Cosmos DB "
    "partition keys. OrbaTech recommended to add SQL RLS.",
    "CI/CD Maturity: Agent Red has a mature 13-phase pipeline. OrbaTech needs to build "
    "this from the ground up.",
]
for item in arch_diffs:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('A.2 Shared Infrastructure Opportunities', level=2)
shared_items = [
    "Shared Azure Key Vault patterns for secrets management",
    "Common Azure Monitor / Application Insights workspace",
    "Shared Azure Communication Services instance for email and SMS",
    "Common CI/CD pipeline templates (GitHub Actions)",
    "Shared Azure Front Door configuration for unified SSL and WAF",
]
for item in shared_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ANNEX B
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Annex B: OrbaTech + Agent Red Integration', level=1)
doc.add_paragraph(
    "This annex describes a bidirectional integration architecture between OrbaTech CRM "
    "and Agent Red AI Customer Service."
)
add_chart(doc, 'fig9',
    "Figure 9: Bidirectional integration architecture.",
    width=Inches(6.5))

doc.add_heading('B.1 CRM Data to Agent Red (Left-to-Right)', level=2)
l2r = [
    "Contact Context: Customer profiles, purchase history, and communication preferences "
    "enable personalized AI responses.",
    "Opportunity Status: Active deals and pipeline stages inform Agent Red's conversation "
    "routing for priority handling.",
    "Activity History: Recent emails, calls, and meetings provide conversation context.",
    "Custom Fields: Tenant-specific custom fields enable Agent Red to adapt to each "
    "business's CRM taxonomy.",
]
for item in l2r:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('B.2 Agent Red to CRM (Right-to-Left)', level=2)
r2l = [
    "Conversation Summaries: AI-generated summaries written to CRM contact records.",
    "Sentiment Analysis: Customer sentiment scores attached to opportunity records.",
    "Quality Metrics: Per-interaction quality scores for sales manager insights.",
    "Lead Qualification: AI-assessed purchase intent signals create/update opportunities.",
]
for item in r2l:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('B.3 Integration Technical Requirements', level=2)
int_reqs = [
    "OAuth 2.0 service-to-service authentication with client credentials flow",
    "REST API endpoints on both platforms with OpenAPI/Swagger documentation",
    "Webhook-based event notifications for real-time data synchronization",
    "Tenant mapping table linking OrbaTech TenantId to Agent Red tenant identifiers",
    "Data transformation layer for field mapping between CRM entities and Agent Red models",
    "Rate limiting and circuit breaker patterns for fault-tolerant integration",
    "Audit logging of all cross-platform data exchanges for compliance",
]
for item in int_reqs:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ANNEX C: SOURCES AND METHODOLOGY (NEW)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Annex C: Sources and Methodology', level=1)
doc.add_paragraph(
    "This annex documents the sources used in this evaluation to provide clear provenance "
    "for all findings and claims."
)
# CORRECTED per Codex finding #5: Explicit source separation

doc.add_heading('C.1 Primary Source: Repository Analysis', level=2)
doc.add_paragraph(
    "The primary source for this evaluation is the OrbaTech GitHub repository "
    "(github.com/QGolle/OrbaTech) as of April 7, 2026. Findings attributed to repository "
    "analysis are derived from direct inspection of checked-in source code, configuration "
    "files, project files, and CI workflow definitions."
)
doc.add_paragraph(
    "Key files examined include: OrbaTech.sln, OrbaTech/OrbaTech.csproj, "
    "OrbaTech.Client/OrbaTech.Client.csproj, Tests/Tests.csproj, OrbaTech/Program.cs, "
    "OrbaTech/Data/ApplicationDbContext.cs, OrbaTech/Services/AttachmentService.cs, "
    "OrbaTech/appsettings.json, OrbaTech/appsettings.Production.json, "
    "EmailService/appsettings.json, .github/workflows/clean-orbatech-csproj.yml, "
    "DB/SP/*.sql, and test files under Tests/Services/, Tests/Data/Repositories/, "
    "and Tests/Components/."
)

doc.add_heading('C.2 Secondary Source: External Market Research', level=2)
doc.add_paragraph(
    "The following sections contain information derived from external market research, "
    "not from the OrbaTech repository:"
)
external_items = [
    "Section 2 (Project Overview): Target markets, website URL",
    "Section 7 (Competitive Landscape): Competitor pricing and market positioning",
    "Section 8 (Feature Comparison Matrix): Competitor capabilities (marked with \u2020)",
    "Section 12 (Azure Cost Projections): Azure pricing estimates from Azure Pricing Calculator",
    "Section 14 (Recommended Technical Priorities): Timeline estimates based on industry benchmarks",
]
for item in external_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('C.3 Revision History', level=2)
add_table(doc, ["Version", "Date", "Changes"], [
    ["1.0", "April 7, 2026", "Initial evaluation report."],
    ["2.0", "April 7, 2026",
     "Revised per Loyal Opposition (Codex) review: (1) Corrected testing assessment to "
     "acknowledge existing test project with xUnit/bUnit/Testcontainers/FsCheck. "
     "(2) Corrected multi-tenant section to acknowledge existing EF Core global query "
     "filters and identify IgnoreQueryFilters() bypass pattern as the actual risk. "
     "(3) Corrected attachment security finding to acknowledge existing path traversal "
     "protection in AttachmentService. (4) Rephrased code review finding to avoid "
     "unsupported inference about PR workflow. (5) Added Annex C to separate repo-backed "
     "findings from external market research. (6) Added Section 8: Detailed Feature "
     "Comparison Matrix covering provisioning, billing, compliance, integrations, "
     "documentation, and architecture. (7) Added two additional competitors (Freshsales, "
     "Monday CRM). (8) Updated Testing score from 2/10 to 4/10."],
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# DISCLAIMER + COPYRIGHT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Disclaimer', level=1)
doc.add_paragraph(
    "This report was prepared by Remaker Digital based on publicly available information in "
    "the OrbaTech GitHub repository as of April 7, 2026. Repository-backed findings reflect "
    "the state of the codebase at that point in time. External market research (competitor "
    "pricing, compliance certifications, integration counts) is sourced from publicly "
    "available vendor documentation and is provided for informational purposes. Azure pricing "
    "estimates are approximate and sourced from the Azure Pricing Calculator."
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    "\u00a9 2026 Remaker Digital, a DBA of VanDusen & Palmeter, LLC. All rights reserved."
)
run.italic = True

# ── Save ─────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Report saved to: {OUTPUT_PATH}")
print(f"Sections: 14 + 3 Annexes")
print(f"Key changes from v1:")
print(f"  - 5 Codex corrections applied")
print(f"  - New Section 8: Detailed Feature Comparison Matrix (6 dimensions, 7 platforms)")
print(f"  - New Annex C: Sources and Methodology")
print(f"  - 2 additional competitors (Freshsales, Monday CRM)")
print(f"  - Testing score: 2/10 -> 4/10")
